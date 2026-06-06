"""
Update thesis with all light-background figures
Maps thesis figure numbers to the newly generated light-background images
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

FIG_DIR = os.path.join(os.path.dirname(__file__), "docs", "figures")
doc = Document("docs/thesis_Edward_Ogbei_PCZ.docx")

def find_para(text):
    text_l = text.lower()
    for i, p in enumerate(doc.paragraphs):
        if text_l in p.text.lower():
            return i
    return -1

def insert_figure_after_text(search_text, image_file, caption_text, width=5.5):
    """Insert figure after paragraph containing search_text."""
    idx = find_para(search_text)
    if idx >= 0:
        ref = doc.paragraphs[idx]
        # Insert image
        img_p = OxmlElement("w:p")
        ref._element.addnext(img_p)
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        img_p.insert(0, pPr)

        tmp = doc.add_paragraph()
        run = tmp.add_run()
        run.add_picture(os.path.join(FIG_DIR, image_file), width=Inches(width))
        tmp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r = tmp._element.find(qn("w:r"))
        if r is not None:
            img_p.append(r)
        tmp._element.getparent().remove(tmp._element)

        # Insert caption
        cap_p = OxmlElement("w:p")
        img_p.addnext(cap_p)
        pPr2 = OxmlElement("w:pPr")
        jc2 = OxmlElement("w:jc")
        jc2.set(qn("w:val"), "center")
        pPr2.append(jc2)
        cap_p.insert(0, pPr2)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        i_el = OxmlElement("w:i")
        sz_el = OxmlElement("w:sz")
        sz_el.set(qn("w:val"), "18")
        rPr.append(i_el)
        rPr.append(sz_el)
        r.append(rPr)
        t_el = OxmlElement("w:t")
        t_el.text = caption_text
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t_el)
        cap_p.append(r)
        return True
    return False

# Figure mapping: Find text to insert after -> (image_file, caption)
figures_to_add = [
    ("The Rubik's Cube is one of those rare objects",
     "fig_01_system_architecture.png",
     "Fig. 1.1. System architecture overview showing the cube engine, solvers, "
     "and visualization components. The AI and classical solvers operate on the same "
     "cubie-based state representation."),

    ("The set of all reachable configurations forms a group under move composition",
     "fig_09_state_space.png",
     "Fig. 3.1. State space and mathematical structure of the Rubik's Cube. "
     "The cube group has 43,252,003,274,489,856,000 distinct states reachable "
     "from the solved configuration."),

    ("For algorithmic purposes, it is common to represent the cube",
     "fig_02_mlp_architecture.png",
     "Fig. 4.1. Cubie-level state representation using corner permutation (cp), "
     "corner orientation (co), edge permutation (ep), and edge orientation (eo). "
     "This is more compact than facelet representation."),

    ("The final production model was trained on a Google Colab",
     "fig_03_encoding.png",
     "Fig. 4.2. One-hot encoding of cube state: each of 54 facelets is encoded "
     "as 6 numbers (0 or 1 for each color), producing a 324-dimensional input vector."),

    ("The RubiksMLP is a simple feedforward network",
     "fig_04_curriculum_flow.png",
     "Fig. 4.3. Curriculum learning flow diagram. The 80 percent solve-rate gate "
     "controls progression to harder depths."),

    ("ReLU was chosen over sigmoid or tanh",
     "fig_02_mlp_architecture.png",
     "Fig. 5.1. MLP architecture detail: three hidden layers (256, 256, 128 units) "
     "with ReLU activations between dense layers."),

    ("Beam search maintains a set of surviving paths",
     "fig_10_beam_search.png",
     "Fig. 5.2. Beam search tree with width 5: at each step, all paths are "
     "expanded and the top-5 are retained by cumulative log-probability score."),

    ("Dataset files are stored as compressed NumPy archives",
     "fig_05_loss_curves.png",
     "Fig. 6.1. Training loss curves by curriculum depth. The spike at each "
     "depth transition represents the model encountering harder scrambles."),

    ("The strategy used here is 80 percent replacement",
     "fig_06_solve_rates.png",
     "Fig. 6.2. Solve rate at each curriculum stage during training. "
     "The horizontal line marks the 80 percent advancement threshold."),

    ("All training code is hardware-agnostic",
     "fig_08_timing.png",
     "Fig. 6.3. Comparative training time: CPU (laptop) requires hours per depth, "
     "while GPU (Colab T4) reduces this to minutes per depth."),

    ("The full web interface consists of a 3D renderer and control panel",
     "fig_7_1_interface.png",
     "Fig. 7.1. Full web interface showing the Three.js 3D cube on the left and "
     "control sidebar on the right. The interface supports interactive scrambling, "
     "animation playback, and solver selection."),

    ("The user can select any of the three solvers",
     "fig_7_2_kociemba_solve.png",
     "Fig. 7.2. Animated Kociemba solve in progress. The solution panel shows the "
     "move sequence and real-time solve statistics during playback."),

    ("Beam search timing scales more steeply with depth",
     "fig_7_3_ai_beam.png",
     "Fig. 7.3. AI beam search solver in action. The interface displays the solution "
     "path and solver statistics including inference time and number of active paths."),

    ("Table 8.1 summarises the solve rates",
     "fig_07_ai_vs_kociemba.png",
     "Fig. 8.1. Benchmark comparison: solve rate (left) and solution time on log scale "
     "(right) for Kociemba versus AI Greedy and AI Beam Search across depths 1 to 10."),

    ("These results support the main thesis claim",
     "real_loss_curves.png",
     "Fig. 8.2. Training loss curves for the three model sizes across all curriculum "
     "depths, showing convergence at each stage and the characteristic spikes at "
     "depth transitions."),

    ("The results demonstrate a clear relationship",
     "real_solve_rate_bars.png",
     "Fig. 8.3. Ablation study results: solve rates for small, medium, and large models "
     "at each depth, showing the relationship between parameter count and curriculum "
     "depth reached."),
]

print("Updating thesis with all light-background figures...\n")
count = 0
for search_text, img_file, caption in figures_to_add:
    # Check if figure file exists
    if os.path.exists(os.path.join(FIG_DIR, img_file)):
        if insert_figure_after_text(search_text, img_file, caption):
            print(f"[OK] Added figure: {img_file}")
            count += 1
        else:
            print(f"[SKIP] Could not find insertion point for {img_file}")
    else:
        print(f"[MISSING] Figure file not found: {img_file}")

doc.save("docs/thesis_Edward_Ogbei_PCZ.docx")
print(f"\n[DONE] Updated thesis with {count} light-background figures.")
print("Now convert to PDF with: python -c \"from docx2pdf import convert; convert(...)\"")
