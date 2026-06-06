from docx import Document

doc = Document("docs/thesis_Edward_Ogbei_PCZ.docx")

search_terms = [
    "cubie representation",
    "RubiksMLP is a simple",
    "Training proceeds in a loop",
    "Beam search maintains",
]

print("Searching for chapter locations...\n")
for term in search_terms:
    print(f"Looking for: '{term}'")
    found = False
    for i, p in enumerate(doc.paragraphs):
        if term.lower() in p.text.lower():
            print(f"  FOUND at paragraph {i}: {p.text[:80]}")
            found = True
    if not found:
        print(f"  NOT FOUND - searching for related keywords...")
        keywords = term.split()
        for keyword in keywords:
            for i, p in enumerate(doc.paragraphs):
                if keyword.lower() in p.text.lower():
                    print(f"    Para {i}: {p.text[:80]}")
                    break
    print()
