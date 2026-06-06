"""
embed_figures_in_thesis.py
--------------------------
Embeds all thesis figures and math equations directly into the thesis docx.

Run: python docs/embed_figures_in_thesis.py
"""

import os, json, copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

FIG  = os.path.join(os.path.dirname(__file__), "figures")
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

doc = Document(os.path.join(ROOT, "docs", "thesis_Edward_Ogbei_PCZ.docx"))


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def find(text):
    tl = text.lower()
    for i, p in enumerate(doc.paragraphs):
        if tl in p.text.lower():
            return i
    return -1


def replace(para, text):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def fig_path(name):
    return os.path.join(FIG, name)


def insert_figure_after(ref_para, img_file, caption_text, width_in=5.8):
    """Insert an image paragraph + caption paragraph after ref_para."""
    # Image paragraph
    img_p = OxmlElement("w:p")
    ref_para._element.addnext(img_p)

    # Build the picture run via python-docx by creating a temp paragraph
    tmp_para = doc.add_paragraph()
    run = tmp_para.add_run()
    run.add_picture(fig_path(img_file), width=Inches(width_in))
    tmp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Move the run XML into img_p
    img_p.append(tmp_para._element.find(qn("w:r")))
    pPr = OxmlElement("w:pPr")
    jc  = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    img_p.insert(0, pPr)
    tmp_para._element.getparent().remove(tmp_para._element)

    # Caption paragraph (italic, centered, 9pt)
    cap_p = OxmlElement("w:p")
    img_p.addnext(cap_p)
    pPr2  = OxmlElement("w:pPr")
    jc2   = OxmlElement("w:jc"); jc2.set(qn("w:val"), "center")
    pPr2.append(jc2)
    cap_p.insert(0, pPr2)
    r     = OxmlElement("w:r")
    rPr   = OxmlElement("w:rPr")
    i_el  = OxmlElement("w:i")
    sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "18")
    rPr.append(i_el); rPr.append(sz_el)
    r.append(rPr)
    t_el  = OxmlElement("w:t")
    t_el.text = caption_text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t_el)
    cap_p.append(r)

    # Return the caption paragraph as the "last inserted" so callers can chain
    return cap_p


def insert_para_after(ref_para, text, bold=False, italic=False, size_pt=None, center=False):
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    pPr = OxmlElement("w:pPr")
    if center:
        jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center")
        pPr.append(jc)
    new_p.insert(0, pPr)
    r   = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:   rPr.append(OxmlElement("w:b"))
    if italic: rPr.append(OxmlElement("w:i"))
    if size_pt:
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size_pt * 2))
        rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    new_p.append(r)
    # Return a lightweight wrapper to allow chaining
    class _Para:
        def __init__(self, el): self._element = el
    return _Para(new_p)


def _el(ref_para):
    """Get the underlying OxmlElement whether ref_para is a docx para or a raw element."""
    return ref_para._element if hasattr(ref_para, "_element") else ref_para


def sp(ref_para):
    """Insert blank spacer paragraph after ref_para."""
    el    = _el(ref_para)
    new_p = OxmlElement("w:p")
    el.addnext(new_p)
    class _P:
        def __init__(self, e): self._element = e
    return _P(new_p)


def insert_math_after(ref_para, eq_file, width_in=3.5):
    """Insert a centred math equation image after ref_para."""
    tmp = doc.add_paragraph()
    run = tmp.add_run()
    run.add_picture(fig_path(eq_file), width=Inches(width_in))
    tmp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    pPr = OxmlElement("w:pPr")
    jc  = OxmlElement("w:jc"); jc.set(qn("w:val"), "center")
    pPr.append(jc)
    new_p.insert(0, pPr)
    r = tmp._element.find(qn("w:r"))
    if r is not None:
        new_p.append(r)
    tmp._element.getparent().remove(tmp._element)

    class _Para:
        def __init__(self, el): self._element = el
    return _Para(new_p)


# ---------------------------------------------------------------------------
# Chapter 3 - Number of cube states equation
# ---------------------------------------------------------------------------
idx = find("The minimum number of face moves required to solve the worst-case")
if idx >= 0:
    ref = doc.paragraphs[idx]
    m   = insert_math_after(ref, "eq_cube_states.png", width_in=5.0)
    sp(m)
    print("Chapter 3: cube states equation inserted.")


# ---------------------------------------------------------------------------
# Chapter 4.4 - One-hot encoding figure + equation
# ---------------------------------------------------------------------------
idx = find("This encoding is simple but carries all the information needed")
if idx >= 0:
    ref = doc.paragraphs[idx]
    m   = insert_math_after(ref, "eq_encoding.png", width_in=5.5)
    cap = insert_figure_after(m, "fig_one_hot.png",
          "Fig. 4.4. One-hot encoding of a single cube facelet. "
          "Each of 54 facelets is encoded as 6 numbers with exactly one 1 "
          "(the colour index) and five 0s, giving a 324-dimensional input vector.",
          width_in=6.0)
    sp(cap)
    print("Chapter 4.4: one-hot figure + equation inserted.")


# ---------------------------------------------------------------------------
# Chapter 5.2 - Architecture figure + ReLU + softmax equations
# ---------------------------------------------------------------------------
idx = find("ReLU was chosen over sigmoid or tanh")
if idx >= 0:
    ref = doc.paragraphs[idx]
    m1  = insert_math_after(ref, "eq_relu.png", width_in=2.8)
    m2  = insert_math_after(m1,  "eq_softmax.png", width_in=4.5)
    cap = insert_figure_after(m2, "fig_architecture.png",
          "Fig. 5.3. RubiksMLP architecture: "
          "324-dimensional one-hot input, three hidden layers (256, 256, 128) "
          "with ReLU activations, and an 18-unit softmax output layer. "
          "Total parameters: 184,210.",
          width_in=6.0)
    sp(cap)
    print("Chapter 5.2: architecture figure + equations inserted.")


# ---------------------------------------------------------------------------
# Chapter 5.3 - Cross-entropy loss equation
# ---------------------------------------------------------------------------
idx = find("The model is trained using the Adam optimiser")
if idx >= 0:
    ref = doc.paragraphs[idx]
    m   = insert_math_after(ref, "eq_loss.png", width_in=3.8)
    sp(m)
    print("Chapter 5/6: cross-entropy loss equation inserted.")


# ---------------------------------------------------------------------------
# Chapter 6.1 - Dataset growth figure
# ---------------------------------------------------------------------------
idx = find("Dataset files are stored as compressed NumPy archives")
if idx >= 0:
    ref = doc.paragraphs[idx]
    cap = insert_figure_after(ref, "fig_dataset_growth.png",
          "Fig. 6.1. Dataset composition and growth at each curriculum depth stage. "
          "Left: stacked bars show new samples (forest green) and retained old "
          "samples (grey). Right: total dataset size grows to over 290,000 "
          "samples by depth 7 for the large model.",
          width_in=6.0)
    sp(cap)
    print("Chapter 6.1: dataset growth figure inserted.")


# ---------------------------------------------------------------------------
# Chapter 6.2 - 80% replacement strategy narrative update
# ---------------------------------------------------------------------------
idx = find("The strategy used here is 80 percent replacement")
if idx >= 0:
    replace(doc.paragraphs[idx],
        "The strategy used here is 80 percent replacement. When advancing from "
        "depth d to depth d+1, the new samples at depth d+1 make up 80 percent "
        "of the combined dataset, while 20 percent is randomly retained from the "
        "previous dataset. This keeps the model reminded of easier cases while "
        "shifting its focus toward harder ones. Without this retention, "
        "the model would forget depth-1 and depth-2 patterns entirely once it "
        "starts training on depth-5 scrambles.")
    print("Chapter 6.2: 80% replacement text updated.")


# ---------------------------------------------------------------------------
# Chapter 6.3 - Curriculum flow figure + threshold equation
# ---------------------------------------------------------------------------
idx = find("The final production model was trained on a Google Colab T4 GPU")
if idx >= 0:
    ref = doc.paragraphs[idx]
    m   = insert_math_after(ref, "eq_threshold.png", width_in=5.5)
    cap = insert_figure_after(m, "fig_curriculum_flow.png",
          "Fig. 6.2. Curriculum learning control flow. "
          "The 80 percent solve-rate gate is the key decision point. "
          "Failing it stops training at that depth, "
          "which is the model's honest capacity ceiling.",
          width_in=5.0)
    sp(cap)
    print("Chapter 6.3: curriculum flow + threshold equation inserted.")


# ---------------------------------------------------------------------------
# Chapter 6.3 - CPU vs GPU figure (why Colab was needed)
# ---------------------------------------------------------------------------
idx = find("All training code is hardware-agnostic")
if idx >= 0:
    ref = doc.paragraphs[idx]
    cap = insert_figure_after(ref, "fig_cpu_vs_gpu.png",
          "Fig. 6.3. Estimated training time per depth stage and cumulative "
          "training time on CPU (laptop) versus GPU T4 (Google Colab). "
          "At depth 6 and 7, the BFS data-generation cost dominates; "
          "the GPU accelerates the neural network training component by "
          "approximately 50 times.",
          width_in=6.0)
    sp(cap)
    print("Chapter 6.3: CPU vs GPU figure inserted.")


# ---------------------------------------------------------------------------
# Chapter 6 NEW SECTION - Per-depth training outcomes
# ---------------------------------------------------------------------------
idx = find("6.4 Implementation Notes")
if idx >= 0:
    ref = doc.paragraphs[idx]

    def ins_before(ref_p, text, bold=False, italic=False):
        new_p = OxmlElement("w:p")
        ref_p._element.addprevious(new_p)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if bold:   rPr.append(OxmlElement("w:b"))
        if italic: rPr.append(OxmlElement("w:i"))
        r.append(rPr)
        t = OxmlElement("w:t"); t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t); new_p.append(r)
        class _P:
            def __init__(self, el): self._element = el
        return _P(new_p)

    # Read ablation data for numbers
    ablation_path = os.path.join(ROOT, "experiments", "plots",
                                 "experiment_results_v2.json")
    with open(ablation_path) as f:
        ablation = json.load(f)
    large = next(r for r in ablation if "large" in r["short"].lower())
    dr    = large["depth_results"]

    new_paras = [
        ("6.4 Training Outcomes - Depth by Depth", True, False),
        ("", False, False),
        ("This section documents what actually happened at each curriculum depth "
         "stage during training of the large model. The per-depth loss curves and "
         "solve rates tell a clear story about how the model progressively learned "
         "harder scrambles.", False, False),
        ("", False, False),
    ]

    # Describe each depth
    for row in dr:
        d    = row["depth"]
        rate = row["solve_rate"]
        n    = row["num_training_samples"]
        if rate >= 0.80:
            verdict = (f"The model cleared the 80 percent gate with a solve rate "
                       f"of {rate:.0%}. Curriculum advanced to depth {d+1}.")
        else:
            verdict = (f"The model achieved {rate:.0%} solve rate, "
                       f"which is below the 80 percent threshold. "
                       f"Curriculum stopped here - this is the model's capacity ceiling.")
        new_paras.append((
            f"Depth {d} ({n:,} training samples): {verdict}",
            False, False))

    new_paras += [
        ("", False, False),
        ("Fig. 6.4 shows the loss curve for each depth stage individually. "
         "The characteristic spike at the start of each panel is the model "
         "encountering harder scrambles for the first time; the subsequent "
         "drop is it adapting. The spike grows larger at each successive depth, "
         "reflecting the increased difficulty. At depth 5, the loss no longer "
         "drops fully, which corresponds to the model failing to clear 80 percent.",
         False, False),
        ("", False, False),
    ]

    for text, bold, italic in reversed(new_paras):
        ins_before(ref, text, bold, italic)

    # Insert the per-depth figure before "6.4 Implementation Notes"
    last_new = ins_before(ref, "")
    # We need to insert the figure inside the doc normally
    tmp_p = doc.add_paragraph()
    run   = tmp_p.add_run()
    run.add_picture(fig_path("fig_depth_by_depth.png"), width=Inches(6.2))
    tmp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fig_xml = OxmlElement("w:p")
    ref._element.addprevious(fig_xml)
    pPr = OxmlElement("w:pPr")
    jc  = OxmlElement("w:jc"); jc.set(qn("w:val"), "center")
    pPr.append(jc); fig_xml.insert(0, pPr)
    r = tmp_p._element.find(qn("w:r"))
    if r is not None:
        fig_xml.append(r)
    tmp_p._element.getparent().remove(tmp_p._element)

    cap_xml = OxmlElement("w:p")
    ref._element.addprevious(cap_xml)
    pPr2 = OxmlElement("w:pPr")
    jc2  = OxmlElement("w:jc"); jc2.set(qn("w:val"), "center")
    pPr2.append(jc2); cap_xml.insert(0, pPr2)
    r2 = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    rPr2.append(OxmlElement("w:i"))
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18"); rPr2.append(sz)
    r2.append(rPr2)
    t2 = OxmlElement("w:t")
    t2.text = ("Fig. 6.4. Training loss at each curriculum depth stage for the "
               "large model (184,210 parameters, 50,000 samples per depth, "
               "100 epochs per depth). Each panel is one depth stage. "
               "The model passed depths 1 to 4 and hit its ceiling at depth 5.")
    t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r2.append(t2); cap_xml.append(r2)

    print("Chapter 6.4: per-depth training outcomes section inserted.")


# ---------------------------------------------------------------------------
# Chapter 5.3 - Beam search figure + equation
# ---------------------------------------------------------------------------
idx = find("Beam search timing scales more steeply with depth")
if idx >= 0:
    ref = doc.paragraphs[idx]
    m   = insert_math_after(ref, "eq_beam.png", width_in=4.8)
    cap = insert_figure_after(m, "fig_beam_search.png",
          "Fig. 5.4. Beam search tree (width=3 for illustration). "
          "At each step all surviving paths are expanded and the "
          "top-3 by cumulative log-probability are kept. "
          "The solved state (green) was found on an alternative branch "
          "that a greedy approach would have discarded.",
          width_in=5.8)
    sp(cap)
    print("Chapter 5/8.3: beam search figure + equation inserted.")


# ---------------------------------------------------------------------------
# Chapter 8.5 - Ablation figures: loss curves, solve rates, model comparison
# ---------------------------------------------------------------------------
idx = find("These results support the main thesis claim: curriculum learning provides")
if idx >= 0:
    ref = doc.paragraphs[idx]

    cap1 = insert_figure_after(ref, "fig_loss_curves.png",
        "Fig. 8.4. Training loss curves across all curriculum depths "
        "for small (amber), medium (blue), and large (green) models. "
        "The depth label in each band shows the solve rate the model "
        "reached before advancing. Loss spikes mark depth transitions.",
        width_in=6.2)
    sp(cap1)

    cap2 = insert_figure_after(cap1, "fig_depth_solve_rates.png",
        "Fig. 8.5. Solve rates at each depth: bars show final performance, "
        "diamonds show the checkpoint rate used by the curriculum gate. "
        "The amber dashed line is the 80 percent advancement threshold.",
        width_in=6.2)
    sp(cap2)

    cap3 = insert_figure_after(cap2, "fig_model_comparison.png",
        "Fig. 8.6. Model capacity versus curriculum depth cleared. "
        "Left: scatter plot of parameter count vs depth reached. "
        "Right: bar chart confirming small clears depth 3, "
        "medium and large both clear depth 4.",
        width_in=6.0)
    sp(cap3)
    print("Chapter 8.5: ablation figures inserted.")


# ---------------------------------------------------------------------------
# Chapter 8.2 - Benchmark figure
# ---------------------------------------------------------------------------
idx = find("Table 8.1 summarises the solve rates")
if idx >= 0:
    ref = doc.paragraphs[idx]
    cap = insert_figure_after(ref, "fig_benchmark.png",
        "Fig. 8.1. Benchmark results: solve rate (left) and solve time in ms "
        "on a log scale (right) for Kociemba, AI Greedy, and AI Beam Search "
        "(width 5) across depths 1 to 10. "
        "50 trials per depth, GPU-trained model.",
        width_in=6.2)
    sp(cap)
    print("Chapter 8.2: benchmark figure inserted.")


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(os.path.join(ROOT, "docs", "thesis_Edward_Ogbei_PCZ.docx"))
print("\nThesis docx saved with all figures embedded.")
