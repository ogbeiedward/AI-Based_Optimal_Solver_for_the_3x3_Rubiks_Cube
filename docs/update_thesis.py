"""
update_thesis.py
----------------
Updates thesis_Edward_Ogbei_PCZ.docx with:
  - New Colab GPU benchmark numbers in Table 8.1
  - Updated results text in sections 8.2, 8.3, 8.4, 9
  - Section 6.3 updated to mention GPU training
  - New Section 8.5: Curriculum Learning Ablation Study
  - Updated abstract
"""

import json
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Load fresh benchmark and ablation data
with open("docs/eval_results.json") as f:
    ev = json.load(f)

with open("experiments/plots/experiment_results_v2.json") as f:
    ablation = json.load(f)

doc = Document("docs/thesis_Edward_Ogbei_PCZ.docx")


def find_para(text):
    text_l = text.lower()
    for i, p in enumerate(doc.paragraphs):
        if text_l in p.text.lower():
            return i
    return -1


def replace_para(para, new_text):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


# ---------------------------------------------------------------------------
# 1. Update Table 8.1
# ---------------------------------------------------------------------------
t = doc.tables[0]
rows_data = [
    ["Depth", "Greedy Rate", "Greedy Avg. Moves",
     "Greedy Avg. Time (ms)", "Beam Rate",
     "Beam Avg. Moves", "Beam Avg. Time (ms)"],
]
for i in range(10):
    g = ev["greedy"][i]
    b = ev["beam_w5"][i]
    rows_data.append([
        str(g["depth"]),
        f"{g['rate']*100:.0f}%",
        f"{g['moves']:.2f}" if g["moves"] else "-",
        f"{g['ms']:.2f}"    if g["ms"]    else "-",
        f"{b['rate']*100:.0f}%",
        f"{b['moves']:.2f}" if b["moves"] else "-",
        f"{b['ms']:.2f}"    if b["ms"]    else "-",
    ])

for row_i, row in enumerate(t.rows):
    for col_i, cell in enumerate(row.cells):
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(rows_data[row_i][col_i])
        if row_i == 0:
            run.bold = True

print("Table 8.1 updated.")


# ---------------------------------------------------------------------------
# 2. Section 8.2 - solve rate summary
# ---------------------------------------------------------------------------
idx = find_para("greedy solver achieves 100% success on scrambles up to depth 3, 97%")
if idx >= 0:
    replace_para(doc.paragraphs[idx],
        "The results confirm that the GPU-trained model solves depths 1 to 3 with "
        "100% accuracy under both strategies. At depth 4 the greedy solver reaches "
        "94% and beam search reaches 98%. Performance declines beyond that: at "
        "depth 5 greedy solves 56% of cases and beam search solves 64%, and at "
        "depth 7 greedy falls to 12% while beam search reaches 22%. Neither solver "
        "finds a solution at depth 10. These figures reflect the model's curriculum "
        "training ceiling: the large model cleared the 80% threshold through depth 4 "
        "and stopped there, which is precisely where reliable performance ends.")
    print(f"Section 8.2 updated (para {idx}).")
else:
    print("WARNING: Section 8.2 paragraph not found.")


# ---------------------------------------------------------------------------
# 3. Section 8.3 - timing
# ---------------------------------------------------------------------------
idx = find_para("Inference times for the greedy solver are remarkably consistent")
if idx >= 0:
    replace_para(doc.paragraphs[idx],
        "Inference times for the greedy solver are consistently fast across all "
        "depths, ranging from 0.62 ms at depth 1 to around 3.75 ms at depth 7. "
        "This near-constant cost reflects the fixed-size forward pass through "
        "the network, which does not depend on scramble complexity. "
        "The Kociemba BFS solver grows dramatically by comparison: from 0.11 ms "
        "at depth 1 to 337.88 ms at depth 7 and over 9,200 ms at depth 10. "
        "At depth 7, the AI greedy solver is roughly 90 times faster than "
        "Kociemba, though with far lower accuracy.")
    print(f"Section 8.3 timing updated (para {idx}).")
else:
    print("WARNING: Section 8.3 paragraph not found.")

idx2 = find_para("Beam search timing scales more steeply with depth")
if idx2 >= 0:
    replace_para(doc.paragraphs[idx2],
        "Beam search timing scales more steeply with depth because each step "
        "expands beam_width candidate paths simultaneously. At depth 3, beam "
        "search takes 4.56 ms, and at depth 7 it reaches 24.80 ms. "
        "Despite this, beam search remains far faster than Kociemba at depths 5 "
        "and above while delivering measurably better accuracy than greedy.")
    print(f"Section 8.3 beam timing updated (para {idx2}).")
else:
    print("WARNING: Section 8.3 beam paragraph not found.")


# ---------------------------------------------------------------------------
# 4. Section 8.4 - discussion
# ---------------------------------------------------------------------------
idx = find_para("The results confirm both hypotheses stated in Chapter 2")
if idx >= 0:
    replace_para(doc.paragraphs[idx],
        "The results confirm both hypotheses set out in Chapter 2. First, the "
        "curriculum-trained model genuinely learns to solve low-depth scrambles "
        "rather than memorising them: it achieves 94% at depth 4 despite never "
        "having seen a specific test scramble before. Second, beam search "
        "consistently improves on greedy search, particularly at depths 4 to 6 "
        "where the model is near its training boundary.")
    print(f"Section 8.4 first paragraph updated (para {idx}).")
else:
    print("WARNING: Section 8.4 paragraph 1 not found.")

idx = find_para("beam search consistently improves on the greedy solver, particularly at the boundary depths")
if idx >= 0:
    replace_para(doc.paragraphs[idx],
        "The most striking result is the speed comparison. At depth 7, the AI "
        "greedy solver responds in under 4 ms while Kociemba takes 338 ms, a "
        "gap of roughly 90 times. At depth 9, Kociemba needs over 4 seconds "
        "while the AI answers in under 1 ms, though it fails to solve the "
        "scramble. This illustrates the fundamental trade-off: the AI trades "
        "accuracy for speed in a way that is predictable and measurable, which "
        "is itself a useful finding.")
    print(f"Section 8.4 second paragraph updated (para {idx}).")
else:
    print("WARNING: Section 8.4 paragraph 2 not found.")


# ---------------------------------------------------------------------------
# 5. Section 6.3 - training hardware
# ---------------------------------------------------------------------------
idx = find_para("Training was conducted entirely on CPU for reproducibility")
if idx >= 0:
    replace_para(doc.paragraphs[idx],
        "The final production model was trained on a Google Colab T4 GPU, which "
        "reduced training time from several hours on a laptop CPU to under "
        "15 minutes. The ablation study comparing three model sizes was run on "
        "the same platform. All training code is hardware-agnostic: it detects "
        "and uses a GPU when available via torch.device, falling back to CPU "
        "otherwise. Batch size was set to 512 for GPU training to make full "
        "use of the available parallelism.")
    print(f"Section 6.3 hardware paragraph updated (para {idx}).")
else:
    print("WARNING: Section 6.3 paragraph not found.")


# ---------------------------------------------------------------------------
# 6. Section 9 - conclusions
# ---------------------------------------------------------------------------
idx = find_para("The experimental results show that the greedy solver achieves 100% success")
if idx >= 0:
    replace_para(doc.paragraphs[idx],
        "The experimental results, measured on a GPU-trained model with 50 trials "
        "per depth, show that the greedy solver achieves 100% success on scrambles "
        "up to depth 3, 94% at depth 4, and 56% at depth 5. Beam search with "
        "width 5 consistently improves on these figures, reaching 98% at depth 4 "
        "and 64% at depth 5. The model was trained with curriculum learning up to "
        "the depth where it could no longer clear the 80% threshold, which for the "
        "large model was depth 4. This represents an honest and measurable capacity "
        "ceiling rather than an arbitrary stopping point.")
    print(f"Section 9 results paragraph updated (para {idx}).")
else:
    print("WARNING: Section 9 paragraph not found.")


# ---------------------------------------------------------------------------
# 7. Abstract
# ---------------------------------------------------------------------------
idx = find_para("The neural network takes a 324-dimensional one-hot encoded input (54 facelets times 6 colours) and predicts the")
if idx >= 0:
    replace_para(doc.paragraphs[idx],
        "The neural network takes a 324-dimensional one-hot encoded input "
        "(54 facelets times 6 colours) and predicts the next move as an 18-class "
        "classification problem. Trained on a T4 GPU with up to 50,000 samples "
        "per curriculum depth, the model achieves 100% solve accuracy on scrambles "
        "up to depth 3, 94% at depth 4, and 56% at depth 5. A beam search "
        "extension with width 5 improves depth-4 accuracy to 98% and depth-5 to "
        "64%. The AI solver is up to 90 times faster than the reference Kociemba "
        "solver at deep scramble depths, trading speed for accuracy in a "
        "predictable and measurable way.")
    print(f"Abstract updated (para {idx}).")
else:
    print("WARNING: Abstract paragraph not found.")


# ---------------------------------------------------------------------------
# 8. Insert Section 8.5 - Ablation Study (before "9. Conclusions")
# ---------------------------------------------------------------------------
def insert_text_before(ref_para, text, bold=False):
    new_p = OxmlElement("w:p")
    ref_para._element.addprevious(new_p)
    r = OxmlElement("w:r")
    if bold:
        rpr = OxmlElement("w:rPr")
        b_el = OxmlElement("w:b")
        rpr.append(b_el)
        r.append(rpr)
    t_el = OxmlElement("w:t")
    t_el.text = text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t_el)
    new_p.append(r)


idx_conc = find_para("9. Conclusions and Future Work")
if idx_conc >= 0:
    ref = doc.paragraphs[idx_conc]

    # Determine what each model achieved
    ablation_summary = []
    for r in ablation:
        dr   = r["depth_results"]
        best = max((d["depth"] for d in dr if d["solve_rate"] >= 0.80), default=0)
        rate_at_next = next((d["solve_rate"] for d in dr if d["depth"] == best + 1), None)
        ablation_summary.append((r["label"], r["n_params"], best, rate_at_next))

    lines = [
        ("8.5 Curriculum Learning Ablation Study", True),
        ("", False),
        ("To understand the relationship between model capacity and curriculum "
         "depth, three model sizes were trained under identical conditions on a "
         "T4 GPU. The small model has 60,434 parameters (hidden layers 128-128), "
         "the medium 118,418 parameters (256-128), and the large 184,210 "
         "parameters (256-256-128). Each used the 80% solve-rate threshold gate "
         "with 15,000, 30,000, and 50,000 samples per depth respectively.", False),
        ("", False),
    ]

    for label, params, best, rate_next in ablation_summary:
        rate_str = f"{rate_next*100:.0f}%" if rate_next is not None else "n/a"
        lines.append((
            f"The {label.split('(')[0].strip().lower()} model ({params:,} parameters) "
            f"cleared the curriculum through depth {best}. At depth {best+1} it "
            f"achieved a solve rate of {rate_str}, which is below the 80% "
            f"threshold, so curriculum stopped there.",
            False
        ))

    lines += [
        ("", False),
        ("The results demonstrate a clear relationship between parameter count and "
         "curriculum depth reached. More parameters allow the model to represent "
         "the more complex move patterns needed at greater scramble depths. "
         "However, even the largest model stalls at depth 4, which confirms that "
         "the current MLP architecture has a practical ceiling for this task. "
         "Reaching depth 6 or beyond would require either a significantly larger "
         "model, a fundamentally different architecture, or a reinforcement "
         "learning approach rather than imitation learning.", False),
        ("", False),
        ("The training loss curves for all three models show a characteristic "
         "spike at each curriculum depth transition followed by recovery. "
         "The larger models recover more completely and to a lower final loss, "
         "which corresponds directly to their higher solve rates at each depth.", False),
        ("", False),
    ]

    for text, bold in reversed(lines):
        insert_text_before(ref, text, bold=bold)

    print(f"Section 8.5 inserted before para {idx_conc}.")
else:
    print("WARNING: Section 9 not found, could not insert 8.5.")


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save("docs/thesis_Edward_Ogbei_PCZ.docx")
print("\nThesis saved: docs/thesis_Edward_Ogbei_PCZ.docx")
