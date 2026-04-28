"""
Generates Project_Explanation.docx with embedded diagrams.
Run from the project root:
    python docs/build_doc.py
"""

import os
import sys
import io
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = os.path.join(os.path.dirname(__file__), "Project_Explanation.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_image(doc, buf, width_inches=5.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))
    if caption:
        add_caption(doc, caption)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Diagram 1: Cube face layout (unfolded net)
# ---------------------------------------------------------------------------

def make_cube_net():
    fig, ax = plt.subplots(figsize=(7, 5.5))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 9)
    ax.set_aspect("equal")
    ax.axis("off")
    fig.patch.set_facecolor("#1a1a2e")
    ax.set_facecolor("#1a1a2e")

    face_colors = {
        "U": "#FFFFFF", "D": "#FFD500",
        "F": "#009E60", "B": "#0051BA",
        "L": "#FF5800", "R": "#C41E3A",
    }
    face_labels = {
        "U": "Up\n(White)", "D": "Down\n(Yellow)",
        "F": "Front\n(Green)", "B": "Back\n(Blue)",
        "L": "Left\n(Orange)", "R": "Right\n(Red)",
    }
    # net layout: (face, col_start, row_start)
    layout = [
        ("U", 3, 6), ("L", 0, 3), ("F", 3, 3),
        ("R", 6, 3), ("B", 9, 3), ("D", 3, 0),
    ]

    for face, cs, rs in layout:
        col = face_colors[face]
        for r in range(3):
            for c in range(3):
                rect = plt.Rectangle(
                    (cs + c, rs + r), 1, 1,
                    facecolor=col,
                    edgecolor="#1a1a2e", linewidth=1.5
                )
                ax.add_patch(rect)
                idx = r * 3 + c
                num_color = "#000000" if face in ("U", "D", "F") else "#ffffff"
                ax.text(cs + c + 0.5, rs + r + 0.5, str(idx),
                        ha="center", va="center", fontsize=7,
                        color=num_color, fontweight="bold")
        cx, cy = cs + 1.5, rs + 1.5
        lbl = face_labels[face]
        ax.text(cx, cy + 0.05, face,
                ha="center", va="center", fontsize=11,
                color="#1a1a2e", fontweight="bold")

    ax.text(6, 8.6,
            "Rubik's Cube Unfolded Net  (54 facelets, numbered 0-8 per face)",
            ha="center", va="center", fontsize=10,
            color="white", fontweight="bold")

    for face, cs, rs in layout:
        lbl = face_labels[face]
        ax.text(cs + 1.5, rs - 0.35, lbl.replace("\n", " "),
                ha="center", va="top", fontsize=7.5, color="#aaaaaa")

    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Diagram 2: One-hot encoding
# ---------------------------------------------------------------------------

def make_onehot_diagram():
    fig, ax = plt.subplots(figsize=(8, 3.2))
    ax.set_facecolor("#1a1a2e")
    fig.patch.set_facecolor("#1a1a2e")
    ax.axis("off")

    colors_hex = ["#FFFFFF", "#C41E3A", "#009E60", "#FFD500", "#FF5800", "#0051BA"]
    color_names = ["White\n(U)", "Red\n(R)", "Green\n(F)", "Yellow\n(D)", "Orange\n(L)", "Blue\n(B)"]
    example_color = 2  # Green is hot

    ax.text(0.5, 0.93, "One-Hot Encoding of a Single Facelet  (example: Green)",
            transform=ax.transAxes, ha="center", fontsize=11,
            color="white", fontweight="bold")

    cell_w = 1.1
    start_x = 0.3
    y_box = 1.5
    y_val = 0.8

    for i, (c, name) in enumerate(zip(colors_hex, color_names)):
        x = start_x + i * cell_w
        edge = "#00ff88" if i == example_color else "#555577"
        lw = 2.5 if i == example_color else 1.0
        rect = plt.Rectangle((x, y_box), 0.9, 0.9,
                              facecolor=c, edgecolor=edge, linewidth=lw)
        ax.add_patch(rect)
        val = "1" if i == example_color else "0"
        txt_col = "#000000" if c in ("#FFFFFF", "#FFD500", "#009E60") else "#ffffff"
        ax.text(x + 0.45, y_box + 0.45, val,
                ha="center", va="center", fontsize=16, fontweight="bold",
                color=txt_col)
        ax.text(x + 0.45, y_val, name,
                ha="center", va="top", fontsize=8, color="#aaaaaa")

    ax.text(start_x - 0.6, y_box + 0.45, "Sticker =",
            ha="right", va="center", fontsize=10, color="white")

    bracket_x = start_x - 0.05
    bracket_end = start_x + 6 * cell_w - 0.15
    ax.annotate("", xy=(bracket_end, y_box - 0.35),
                xytext=(bracket_x, y_box - 0.35),
                arrowprops=dict(arrowstyle="<->", color="#00ff88", lw=1.5))
    ax.text((bracket_x + bracket_end) / 2, y_box - 0.6,
            "6 values per facelet   x   54 facelets   =   324-dim input vector",
            ha="center", va="top", fontsize=9, color="#00ff88")

    ax.set_xlim(-0.2, 8)
    ax.set_ylim(0, 3.2)
    ax.set_aspect("equal")
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Diagram 3: Cubie-level state arrays
# ---------------------------------------------------------------------------

def make_state_arrays_diagram():
    fig, axes = plt.subplots(1, 4, figsize=(10, 2.8))
    fig.patch.set_facecolor("#1a1a2e")

    array_data = [
        ("cp\nCorner Permutation", [0, 1, 2, 3, 4, 5, 6, 7],
         "#C41E3A", "Which corner\nis in each slot"),
        ("co\nCorner Orientation", [0, 0, 0, 0, 0, 0, 0, 0],
         "#FF5800", "How each corner\nis twisted (0/1/2)"),
        ("ep\nEdge Permutation", [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11],
         "#009E60", "Which edge\nis in each slot"),
        ("eo\nEdge Orientation", [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0],
         "#0051BA", "Whether each\nedge is flipped"),
    ]

    for ax, (title, vals, col, desc) in zip(axes, array_data):
        ax.set_facecolor("#1a1a2e")
        ax.axis("off")
        ax.text(0.5, 1.0, title, transform=ax.transAxes,
                ha="center", va="bottom", fontsize=8.5, color="white",
                fontweight="bold", multialignment="center")

        n = len(vals)
        for i, v in enumerate(vals):
            rect = plt.Rectangle((i * 1.05, 0), 0.9, 0.9,
                                  facecolor=col + "55", edgecolor=col,
                                  linewidth=1.2)
            ax.add_patch(rect)
            ax.text(i * 1.05 + 0.45, 0.45, str(v),
                    ha="center", va="center", fontsize=9,
                    color="white", fontweight="bold")
            ax.text(i * 1.05 + 0.45, -0.15, str(i),
                    ha="center", va="top", fontsize=6.5, color="#888888")

        ax.text(n * 1.05 / 2, -0.55, desc,
                ha="center", va="top", fontsize=7.5, color="#aaaaaa",
                multialignment="center")
        ax.set_xlim(-0.1, n * 1.05)
        ax.set_ylim(-1.1, 1.3)

    fig.suptitle("Solved Cube State: Four Arrays (all identity/zeros = solved)",
                 color="white", fontsize=10, y=1.02)
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Diagram 4: Neural network architecture
# ---------------------------------------------------------------------------

def make_nn_diagram():
    fig, ax = plt.subplots(figsize=(9, 4.5))
    ax.set_facecolor("#1a1a2e")
    fig.patch.set_facecolor("#1a1a2e")
    ax.axis("off")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)

    layers = [
        (1.0, "Input\n324", "#555577", 324),
        (3.0, "Dense\n256\n(ReLU)", "#C41E3A", 256),
        (5.0, "Dense\n256\n(ReLU)", "#C41E3A", 256),
        (7.0, "Dense\n128\n(ReLU)", "#FF5800", 128),
        (9.0, "Output\n18\n(Softmax)", "#009E60", 18),
    ]

    n_vis = 7
    for x, label, col, n in layers:
        drawn = min(n_vis, n)
        step = 3.6 / drawn
        for i in range(drawn):
            y = 0.7 + i * step
            circle = plt.Circle((x, y), 0.18, color=col, zorder=3)
            ax.add_patch(circle)
        if n > n_vis:
            ax.text(x, 0.7 + drawn * step + 0.05, "...",
                    ha="center", va="bottom", color="#aaaaaa", fontsize=12)
        ax.text(x, 4.6, label, ha="center", va="bottom",
                fontsize=8.5, color="white", multialignment="center",
                fontweight="bold")

    for i in range(len(layers) - 1):
        x1 = layers[i][0] + 0.2
        x2 = layers[i + 1][0] - 0.2
        drawn1 = min(n_vis, layers[i][3])
        drawn2 = min(n_vis, layers[i + 1][3])
        step1 = 3.6 / drawn1
        step2 = 3.6 / drawn2
        for a in range(drawn1):
            for b in range(drawn2):
                y1 = 0.7 + a * step1
                y2 = 0.7 + b * step2
                ax.plot([x1, x2], [y1, y2], color="#334466", lw=0.3, zorder=1)

    param_labels = ["", "83,200", "65,792", "32,896", "2,322"]
    for i, (x, _, _, _) in enumerate(layers):
        if param_labels[i]:
            ax.text(x, 0.1, param_labels[i] + " params",
                    ha="center", va="bottom", fontsize=7, color="#888888")

    ax.text(5, 0.05, "Total: ~184,000 trainable parameters",
            ha="center", fontsize=9, color="#00ff88")

    ax.set_title("RubiksMLP Architecture", color="white", fontsize=11,
                 fontweight="bold", pad=6)
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Diagram 5: Bidirectional BFS
# ---------------------------------------------------------------------------

def make_bfs_diagram():
    fig, axes = plt.subplots(1, 2, figsize=(10, 4))
    fig.patch.set_facecolor("#1a1a2e")

    titles = ["Standard BFS\n(from scrambled state only)",
              "Bidirectional BFS\n(from both ends, meet in middle)"]
    bgs = ["#1a1a2e", "#1a1a2e"]

    for ax, title in zip(axes, titles):
        ax.set_facecolor("#1a1a2e")
        ax.axis("off")
        ax.set_xlim(-1, 1)
        ax.set_ylim(-0.2, 4.2)
        ax.set_title(title, color="white", fontsize=9.5, fontweight="bold")

    # Left: standard BFS tree
    ax = axes[0]
    depths_nodes = [(0, [0.0]), (1, [-0.6, 0.6]),
                    (2, [-0.9, -0.3, 0.3, 0.9]),
                    (3, [-0.95, -0.65, -0.35, -0.05, 0.25, 0.55, 0.85, 0.95])]
    y_levels = [3.8, 2.7, 1.6, 0.5]
    col = "#C41E3A"
    for lvl, (d, nodes) in enumerate(depths_nodes):
        y = y_levels[lvl]
        for x in nodes:
            c = plt.Circle((x, y), 0.09, color=col, zorder=3)
            ax.add_patch(c)
        if lvl > 0:
            prev_y = y_levels[lvl - 1]
            prev_nodes = depths_nodes[lvl - 1][1]
            for px in prev_nodes:
                for cx in nodes:
                    ax.plot([px, cx], [prev_y - 0.09, y + 0.09],
                            color="#334466", lw=0.5, zorder=1)
    ax.text(0, 4.1, "Scrambled\nstate", ha="center", va="bottom",
            fontsize=7.5, color="#C41E3A")
    ax.text(0, 0.05, "18^10 ~ 3.5 trillion nodes\n(depth 10)", ha="center",
            va="top", fontsize=8, color="#ff8888")

    # Right: bidirectional
    ax = axes[1]
    top_nodes = [(0, [0.0]), (1, [-0.4, 0.4]), (2, [-0.65, -0.15, 0.15, 0.65])]
    bot_nodes = [(0, [0.0]), (1, [-0.4, 0.4]), (2, [-0.65, -0.15, 0.15, 0.65])]
    top_y = [3.8, 2.85, 1.9]
    bot_y = [0.2, 1.15, 2.1]

    for lvl, (d, nodes) in enumerate(top_nodes):
        y = top_y[lvl]
        for x in nodes:
            c = plt.Circle((x, y), 0.09, color="#C41E3A", zorder=3)
            ax.add_patch(c)
        if lvl > 0:
            py = top_y[lvl - 1]
            pn = top_nodes[lvl - 1][1]
            for px in pn:
                for cx in nodes:
                    ax.plot([px, cx], [py - 0.09, y + 0.09],
                            color="#553333", lw=0.5, zorder=1)

    for lvl, (d, nodes) in enumerate(bot_nodes):
        y = bot_y[lvl]
        for x in nodes:
            c = plt.Circle((x, y), 0.09, color="#009E60", zorder=3)
            ax.add_patch(c)
        if lvl > 0:
            py = bot_y[lvl - 1]
            pn = bot_nodes[lvl - 1][1]
            for px in pn:
                for cx in nodes:
                    ax.plot([px, cx], [py + 0.09, y - 0.09],
                            color="#224433", lw=0.5, zorder=1)

    ax.text(0, 4.1, "Scrambled", ha="center", va="bottom",
            fontsize=7.5, color="#C41E3A")
    ax.text(0, -0.05, "Solved", ha="center", va="top",
            fontsize=7.5, color="#009E60")

    mid_y = (top_y[-1] + bot_y[-1]) / 2
    ax.annotate("", xy=(0, top_y[-1] - 0.1), xytext=(0, bot_y[-1] + 0.1),
                arrowprops=dict(arrowstyle="<->", color="#00ff88", lw=1.5))
    ax.text(0.72, mid_y, "Meet\nin\nmiddle", ha="center", va="center",
            fontsize=8, color="#00ff88",
            bbox=dict(boxstyle="round,pad=0.3", facecolor="#224433",
                      edgecolor="#00ff88"))
    ax.text(0, 0.05, "2 x 18^5 ~ 3.4 million nodes\n(~1,000,000x fewer)",
            ha="center", va="top", fontsize=8, color="#88ff88",
            transform=ax.transData)

    fig.suptitle("Why Bidirectional BFS is Faster", color="white",
                 fontsize=11, fontweight="bold", y=1.02)
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Diagram 6: Curriculum learning
# ---------------------------------------------------------------------------

def make_curriculum_diagram():
    fig, ax = plt.subplots(figsize=(9, 3.8))
    ax.set_facecolor("#1a1a2e")
    fig.patch.set_facecolor("#1a1a2e")
    ax.axis("off")
    ax.set_xlim(-0.5, 8.5)
    ax.set_ylim(-0.5, 3.5)

    stages = [
        ("Depth 1\n(1 move from\nsolved)", 0.7, "#C41E3A", "100%"),
        ("Depth 2\n(2 moves)", 2.2, "#FF5800", "100%"),
        ("Depth 3\n(3 moves)", 3.7, "#FFD500", "100%"),
        ("Depth 4\n(4 moves)", 5.2, "#009E60", "100%"),
        ("Depth 5\n(5 moves)", 6.7, "#0051BA", "71%"),
    ]

    for label, x, col, rate in stages:
        rect = FancyBboxPatch((x - 0.55, 1.2), 1.1, 1.4,
                               boxstyle="round,pad=0.08",
                               facecolor=col + "44", edgecolor=col,
                               linewidth=1.8)
        ax.add_patch(rect)
        ax.text(x, 1.95, label, ha="center", va="center",
                fontsize=8, color="white", multialignment="center")
        ax.text(x, 1.0, f"Solve rate: {rate}", ha="center", va="top",
                fontsize=8, color=col, fontweight="bold")

    for i in range(len(stages) - 1):
        x1 = stages[i][1] + 0.55
        x2 = stages[i + 1][1] - 0.55
        ax.annotate("", xy=(x2, 1.9), xytext=(x1, 1.9),
                    arrowprops=dict(arrowstyle="->", color="white", lw=1.5))
        ax.text((x1 + x2) / 2, 2.08, ">80%\nthreshold",
                ha="center", va="bottom", fontsize=6.5, color="#aaaaaa",
                multialignment="center")

    ax.text(3.7, 0.25,
            "At each step: generate new samples, train 50 epochs, evaluate.\n"
            "Only advance if solve rate exceeds 80%. Keep 20% of old data to avoid forgetting.",
            ha="center", va="center", fontsize=8.5, color="#aaaaaa",
            multialignment="center")

    ax.set_title("Curriculum Learning Strategy", color="white",
                 fontsize=11, fontweight="bold", pad=6)
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Diagram 7: System architecture
# ---------------------------------------------------------------------------

def make_architecture_diagram():
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_facecolor("#1a1a2e")
    fig.patch.set_facecolor("#1a1a2e")
    ax.axis("off")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)

    boxes = [
        (5.0, 5.2, "User (Browser / Keyboard)", "#334466", "#aabbdd", "#334466", 4.5, 0.5),
        (5.0, 4.2, "main.py  /  visualization/server.py\n(Entry point + HTTP API)", "#1a3a6e", "#ffffff", "#1a3a6e", 5.5, 0.6),
        (5.0, 3.1, "core/cube.py  (CubieCube state engine)\ncp, co, ep, eo arrays", "#3a1a6e", "#ccaaff", "#9977cc", 5.5, 0.6),
        (2.0, 1.8, "solvers/kociemba_solver.py\nBidirectional BFS", "#1a4a2e", "#88ffaa", "#44aa77", 3.5, 0.7),
        (8.0, 1.8, "solvers/ai_solver.py\nMLP Neural Network", "#4a1a1e", "#ffaaaa", "#aa4444", 3.5, 0.7),
        (5.0, 0.6, "visualization/index.html\nThree.js 3D Web Interface", "#1a2e4a", "#aaccff", "#4477aa", 5.5, 0.6),
    ]

    rects = {}
    for cx, cy, label, fc, tc, ec, w, h in boxes:
        rect = FancyBboxPatch((cx - w / 2, cy - h / 2), w, h,
                               boxstyle="round,pad=0.1",
                               facecolor=fc, edgecolor=ec,
                               linewidth=1.5)
        ax.add_patch(rect)
        ax.text(cx, cy, label, ha="center", va="center",
                fontsize=8, color=tc, multialignment="center",
                fontweight="bold")
        rects[label[:8]] = (cx, cy, h)  # noqa

    arrows = [
        ((5.0, 4.95), (5.0, 4.5), "HTTP requests"),
        ((5.0, 3.9), (5.0, 3.4), "apply_move / is_solved"),
        ((3.75, 3.1), (3.75, 2.15), "cube state"),
        ((6.25, 3.1), (6.25, 2.15), "encoded state\n(324-vector)"),
        ((5.0, 0.9), (5.0, 2.78), "solution moves", True),
    ]

    for (x1, y1), (x2, y2), lbl, *rev in arrows:
        style = "<-" if rev else "->"
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle=style, color="#888888",
                                   lw=1.2))
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(mx + 0.15, my, lbl, ha="left", va="center",
                fontsize=7, color="#aaaaaa", multialignment="center")

    ax.annotate("", xy=(5.0, 0.9), xytext=(5.0, 1.44),
                arrowprops=dict(arrowstyle="->", color="#888888", lw=1.2))

    ax.set_title("System Architecture Overview", color="white",
                 fontsize=11, fontweight="bold", pad=6)
    return fig_to_bytes(fig)


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("AI-Based Optimal Solver for the 3x3 Rubik's Cube", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Full Project Explanation")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].italic = True
    meta = doc.add_paragraph("Edward Ogbei   |   M.Sc. Artificial Intelligence   |   PCZ")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    # -------------------------------------------------------------------------
    add_heading(doc, "1. What the Project Does", 1)
    add_body(doc,
        "This project builds a complete system for simulating, scrambling, and solving a "
        "standard 3x3 Rubik's Cube. Two solving methods are implemented and compared: a "
        "classical algorithm based on bidirectional BFS (inspired by the Kociemba two-phase "
        "approach), and a trained neural network (Multi-Layer Perceptron). Both solvers are "
        "evaluated in a benchmarking suite, and the entire system is accessible through an "
        "interactive 3D web interface where the user can manipulate the cube and watch either "
        "solver animate its solution in real time.")
    add_body(doc,
        "The full pipeline is covered: mathematical cube model, move engine, state encoding, "
        "dataset generation, machine learning training, evaluation, and visualisation.")

    # System architecture diagram
    add_image(doc, make_architecture_diagram(), 5.8,
              "Figure 1: System architecture. The browser communicates with the Python HTTP "
              "backend, which drives the cube engine and calls either the Kociemba or AI solver.")

    # -------------------------------------------------------------------------
    add_heading(doc, "2. Why the Rubik's Cube is a Hard Problem", 1)
    add_body(doc,
        "A standard 3x3 Rubik's Cube has exactly 43,252,003,274,489,856,000 legal "
        "configurations (roughly 43 quintillion). Despite this, it has been mathematically "
        "proven that every configuration can be solved in at most 20 moves. That maximum "
        "is called God's Number.")
    add_body(doc,
        "The state space is far too large to search exhaustively. Even a computer checking "
        "one million states per second would take over a billion years to visit all of them. "
        "This is why smart algorithms and learned heuristics are necessary.")
    add_body(doc,
        "Additionally, not all piece arrangements are physically reachable. Three mathematical "
        "constraints must always hold for a state to be valid:")
    doc.add_paragraph("Corner orientations must sum to a multiple of 3.", style="List Bullet")
    doc.add_paragraph("Edge orientations must sum to a multiple of 2.", style="List Bullet")
    doc.add_paragraph(
        "The permutation parity of corners must equal the permutation parity of edges.",
        style="List Bullet")
    add_body(doc,
        "These constraints are enforced throughout the project via an is_legal() check on "
        "the cube object.")

    # Cube net diagram
    add_image(doc, make_cube_net(), 5.8,
              "Figure 2: Unfolded net of the Rubik's Cube showing all 54 facelets. "
              "Each face has 9 stickers numbered 0 to 8. Six faces give 54 stickers total.")

    # -------------------------------------------------------------------------
    add_heading(doc, "3. The Cube Simulator", 1)
    add_heading(doc, "File: core/cube.py, core/cubie.py", 2)

    add_body(doc,
        "The cube simulator is a mathematical model stored in memory as four arrays. "
        "It does not store colours directly. Instead it stores the positions and orientations "
        "of the physical pieces (called cubies).")
    add_heading(doc, "Why cubie-level representation?", 2)
    add_body(doc,
        "There are two common ways to model the cube. The naive approach stores 54 coloured "
        "stickers in a flat array. This is simple, but moves become complex transformations "
        "and it is easy to accidentally create illegal states.")
    add_body(doc,
        "The cubie model used here stores the actual pieces and tracks where each piece is "
        "and how it is twisted. This is the standard used in academic and competitive solving "
        "research because it directly encodes the group-theoretical structure of the cube, "
        "makes legality checking trivial, and interfaces naturally with the Kociemba algorithm.")

    add_table(doc,
        ["Array", "Size", "What it stores"],
        [
            ("cp  (corner permutation)", "8 integers", "Which corner piece is in each of the 8 corner slots"),
            ("co  (corner orientation)", "8 integers", "How each corner is twisted: 0, 1, or 2 (representing 0, 120, 240 degrees)"),
            ("ep  (edge permutation)",   "12 integers", "Which edge piece is in each of the 12 edge slots"),
            ("eo  (edge orientation)",   "12 integers", "Whether each edge is flipped: 0 = correct, 1 = flipped"),
        ]
    )

    add_image(doc, make_state_arrays_diagram(), 5.8,
              "Figure 3: The four arrays representing a solved cube. All corner and edge "
              "pieces are in their home positions (cp/ep = identity) with zero twist/flip (co/eo = 0).")

    add_body(doc,
        "The eight corners are named by the three faces they touch: URF, UFL, ULB, UBR, "
        "DFR, DLF, DBL, DRB. The twelve edges are named by two faces: UR, UF, UL, UB, "
        "DR, DF, DL, DB, FR, FL, BL, BR. This naming convention is standard in the Kociemba "
        "algorithm literature.")

    # -------------------------------------------------------------------------
    add_heading(doc, "4. The Move Engine", 1)
    add_heading(doc, "File: core/moves.py", 2)

    add_body(doc,
        "The move engine defines all 18 standard Rubik's Cube moves and how they transform "
        "the cube state. Every move is described as a set of permutation cycles applied to "
        "the four state arrays.")
    add_heading(doc, "Why 18 moves?", 2)
    add_body(doc,
        "The six faces (U, D, R, L, F, B) each have three variants: a clockwise quarter "
        "turn (R), a counter-clockwise quarter turn (R'), and a half turn (R2). That gives "
        "6 x 3 = 18 moves. These 18 moves form a generating set for the entire Rubik's Cube "
        "group, meaning any legal state can be reached from any other state using only these "
        "moves.")
    add_body(doc,
        "Every move is stored as a MoveDefinition object with corner permutation cycles, "
        "corner orientation deltas, edge permutation cycles, and edge orientation deltas. "
        "The U and D moves do not change orientations because turning the top or bottom face "
        "does not physically twist corners or flip edges. The F, B, R, and L moves do affect "
        "orientations.")

    # -------------------------------------------------------------------------
    add_heading(doc, "5. State Encoding (One-Hot Vectors)", 1)
    add_heading(doc, "File: core/state_encoder.py", 2)

    add_body(doc,
        "The neural network cannot accept arrays of positions and orientations directly as "
        "input. It needs a fixed-size numerical vector. The state encoder converts the cube "
        "into a 324-dimensional vector using one-hot encoding.")
    add_heading(doc, "What one-hot encoding means", 2)
    add_body(doc,
        "The cube has 54 visible sticker positions (9 per face, 6 faces). Each sticker shows "
        "one of six colours. One-hot encoding represents each colour as a vector of six binary "
        "values where exactly one is 1 and the rest are 0. For all 54 stickers, this produces "
        "54 x 6 = 324 values. This 324-element vector is the input to the neural network.")

    add_image(doc, make_onehot_diagram(), 5.5,
              "Figure 4: One-hot encoding of a single facelet. Green is represented as "
              "[0,0,1,0,0,0]. All 54 facelets concatenated give the 324-dim input vector.")

    add_heading(doc, "Why one-hot instead of integers?", 2)
    add_body(doc,
        "Encoding colours as integers (White=0, Red=1, etc.) would imply an ordering "
        "relationship that does not exist. White is not less than Red in any meaningful way. "
        "One-hot encoding avoids this by giving each colour an independent dimension, allowing "
        "the network to learn separate weights for each colour at each position.")

    # -------------------------------------------------------------------------
    add_heading(doc, "6. The Dataset Generator", 1)
    add_heading(doc, "Files: data/dataset_generator.py, data/dataset_stats.py", 2)

    add_heading(doc, "The core idea: expert cloning", 2)
    add_body(doc,
        "The training data is generated by asking the Kociemba solver what the best next "
        "move is from any given scrambled state. The neural network then learns to imitate "
        "the expert. This technique is called behavioural cloning or imitation learning.")
    add_body(doc,
        "For each training sample: a cube is scrambled by N random moves, the Kociemba "
        "solver finds the full solution, and the first move of that solution is recorded as "
        "the correct label. The dataset is a collection of (encoded state, correct next move) "
        "pairs.")
    add_heading(doc, "Curriculum learning", 2)
    add_body(doc,
        "Training on random 20-move scrambles from scratch does not work well. The problem "
        "is too hard initially. Instead, training follows a curriculum: start with depth 1 "
        "(one move from solved), then advance only when performance at the current depth "
        "exceeds 80%.")

    add_image(doc, make_curriculum_diagram(), 5.8,
              "Figure 5: Curriculum learning strategy. Training advances from depth 1 to "
              "depth 5 only when the solve rate at each level exceeds 80%.")

    add_heading(doc, "The 80% replacement strategy", 2)
    add_body(doc,
        "When advancing from depth N to N+1, the dataset is updated: 80% is freshly "
        "generated at the new depth, and 20% is kept from the previous dataset. This "
        "prevents the model from forgetting how to solve easy states while learning harder "
        "ones. The phenomenon of forgetting previously learned skills is called catastrophic "
        "forgetting, and the replacement strategy mitigates it.")
    add_body(doc,
        "The saved dataset files (data/datasets/dataset_depth_1.npz through "
        "dataset_depth_8.npz) contain the cumulative training data at each stage.")

    # -------------------------------------------------------------------------
    add_heading(doc, "7. The Kociemba Solver (Classical Algorithm)", 1)
    add_heading(doc, "Files: solvers/kociemba_solver.py, solvers/bfs_solver.py", 2)

    add_heading(doc, "Bidirectional BFS", 2)
    add_body(doc,
        "The classical solver uses bidirectional BFS (Breadth-First Search) to find the "
        "shortest solution. Standard BFS from the scrambled state grows at 18^N nodes. At "
        "depth 10, that is 18^10 = approximately 3.5 trillion nodes.")
    add_body(doc,
        "Bidirectional BFS searches simultaneously from both the scrambled state and the "
        "solved state, and stops when the two frontiers meet. This reduces the tree size "
        "to 2 x 18^(N/2). At depth 10 that is approximately 3.4 million nodes, roughly "
        "one million times fewer.")

    add_image(doc, make_bfs_diagram(), 5.8,
              "Figure 6: Bidirectional BFS vs standard BFS. Searching from both ends "
              "and meeting in the middle reduces the search space by roughly 10^6 at depth 10.")

    add_heading(doc, "Implementation details", 2)
    add_body(doc,
        "The cube state is stored as a plain Python tuple of 40 integers (8 cp + 8 co + "
        "12 ep + 12 eo). Tuples are hashable and can be stored in sets, making frontier "
        "lookups O(1) instead of O(N). Move pruning avoids redundant sequences: no two "
        "consecutive moves on the same face, and opposite-face pairs are always applied "
        "in canonical order.")
    add_body(doc,
        "When the interactive web viewer is used and the scramble history is available, "
        "the solver skips BFS entirely and reverses the history. If R, U, F were applied, "
        "the solution is F', U', R'. This is instantaneous.")

    # -------------------------------------------------------------------------
    add_heading(doc, "8. The AI Solver (Neural Network)", 1)
    add_heading(doc, "File: solvers/ai_solver.py", 2)

    add_heading(doc, "Architecture", 2)
    add_body(doc,
        "The neural network is a fully-connected feedforward network (Multi-Layer Perceptron) "
        "with four layers: an input of 324 units, two hidden layers of 256 units, one hidden "
        "layer of 128 units, and an output layer of 18 units (one per move). Total trainable "
        "parameters: approximately 184,000.")

    add_image(doc, make_nn_diagram(), 5.8,
              "Figure 7: RubiksMLP architecture. The 324-dim one-hot state vector passes "
              "through three ReLU layers before producing 18 move-probability scores.")

    add_heading(doc, "Why an MLP and not a CNN or RNN?", 2)
    add_body(doc,
        "A Convolutional Neural Network (CNN) is designed for data with spatial locality, "
        "where nearby elements share structure (like image pixels). The cube state encoded "
        "as a flat 324-vector does not have this property. Adjacent stickers on the physical "
        "cube are often far apart in the flat vector.")
    add_body(doc,
        "A Recurrent Neural Network (RNN) is designed for sequential data where past context "
        "matters (like text). A single cube state is a fixed snapshot, not a sequence. An "
        "RNN would add unnecessary complexity.")
    add_body(doc,
        "An MLP takes the full state vector and learns arbitrary non-linear mappings from it "
        "to move probabilities. With three hidden layers it can represent the complex patterns "
        "in cube states while remaining trainable in hours on a standard GPU.")

    add_heading(doc, "Design choices", 2)
    add_table(doc,
        ["Choice", "What it is", "Why it was chosen"],
        [
            ("ReLU activation", "f(x) = max(0, x)", "Cheap to compute, avoids vanishing gradients, enables fast learning"),
            ("Cross-Entropy Loss", "Standard classification loss", "The task is 18-class classification: which move to apply next"),
            ("Adam optimiser", "Adaptive per-parameter learning rate", "Converges faster than plain SGD, less manual tuning, default lr=0.001"),
            ("Batch size 64", "64 samples per gradient update", "Balances training speed and gradient quality"),
            ("50 epochs per depth", "Full passes over the training set", "Enough for convergence at each curriculum level"),
        ]
    )

    add_heading(doc, "Greedy inference", 2)
    add_body(doc,
        "At inference time the model solves a cube greedily: encode the state, pass through "
        "the network, apply softmax to get probabilities, pick the highest-probability move, "
        "apply it, and repeat until solved or 50 steps have been attempted. This achieves "
        "100% solve rate at depths 1 to 4 and approximately 71% at depth 5.")

    # -------------------------------------------------------------------------
    add_heading(doc, "9. The 3D Visualisation System", 1)
    add_heading(doc, "Files: visualization/server.py, visualization/index.html", 2)

    add_body(doc,
        "The project has two visualisers. The web-based 3D viewer using Three.js (WebGL) "
        "is the primary demo interface. The desktop viewer using matplotlib was the original "
        "prototype built during early development.")
    add_heading(doc, "The Python HTTP backend", 2)
    add_body(doc,
        "The backend is a plain Python HTTP server (no Flask or Django) listening on port "
        "8080. It exposes a small REST API. No external web framework was added because the "
        "API is small enough that Python's built-in http.server module handles it.")

    add_table(doc,
        ["Endpoint", "Method", "What it does"],
        [
            ("/", "GET", "Serve the index.html file"),
            ("/state", "GET", "Return current cube state as JSON"),
            ("/move", "POST", "Apply one move (e.g. R)"),
            ("/scramble", "POST", "Apply a random scramble"),
            ("/reset", "POST", "Reset to solved state"),
            ("/solve", "POST", "Solve with Kociemba, return solution"),
            ("/solve_ai", "POST", "Solve with neural network, return solution"),
        ]
    )

    add_heading(doc, "The Three.js frontend", 2)
    add_body(doc,
        "Three.js is a JavaScript library for 3D graphics in the browser using WebGL. It "
        "was chosen because it runs in any modern browser without plugins, produces smooth "
        "60 fps rendering, and has built-in support for 3D objects, lighting, and camera "
        "controls. Raw WebGL would require hundreds of lines of shader code for the same result.")
    add_body(doc,
        "The cube is built as 27 individual cubie objects (one per 3x3x3 position). Each "
        "cubie has a dark internal box and coloured sticker faces. After each move, syncState() "
        "uses the 24-element rotation group of the cube to compute the correct 3D quaternion "
        "for every piece and smoothly animates it into place.")

    # -------------------------------------------------------------------------
    add_heading(doc, "10. Benchmarking and Experiments", 1)
    add_heading(doc, "Files: experiments/ai_vs_kociemba.py, experiments/kociemba_benchmark.py", 2)

    add_body(doc,
        "ai_vs_kociemba.py runs both solvers on the same randomly scrambled cubes across "
        "depths 1 to 5 and prints a side-by-side comparison of solve rate, average number "
        "of moves, and average wall-clock time. It also reports AI confidence distributions "
        "and failure analysis by depth.")
    add_body(doc,
        "kociemba_benchmark.py runs the Kociemba solver alone on a large batch (default 100 "
        "scrambles) and computes mean, median, min, max, and standard deviation for both "
        "time and solution length.")

    add_table(doc,
        ["Depth", "AI Solve Rate", "Kociemba Solve Rate", "Notes"],
        [
            ("1", "100%", "100%", "Trivially easy for both"),
            ("2", "100%", "100%", ""),
            ("3", "100%", "100%", ""),
            ("4", "100%", "~100%", "AI matches expert perfectly"),
            ("5", "~71%", "~100%", "AI starts to struggle; greedy search limitation"),
        ]
    )

    # -------------------------------------------------------------------------
    add_heading(doc, "11. Libraries and Tools Used", 1)

    add_table(doc,
        ["Library / Tool", "Why it was chosen"],
        [
            ("Python 3.14", "Dominant language for ML; richest ecosystem of relevant libraries"),
            ("PyTorch", "Best research deep learning framework; dynamic graphs make debugging easy; runs on CPU and GPU"),
            ("NumPy", "Fast array operations for dataset storage, one-hot encoding, and statistics"),
            ("Matplotlib", "Desktop 3D viewer prototype; no extra install beyond pip"),
            ("Three.js (JS)", "Browser 3D rendering without raw WebGL shader code; 60 fps in any browser"),
            ("pytest", "Standard Python testing; minimal boilerplate; excellent output formatting"),
            ("Docker", "Containerises the app for identical deployment on any machine or cloud host"),
            ("Render.com", "Free cloud hosting used to deploy the live demo"),
            ("Google Colab", "Free GPU environment for training the model (10-20x faster than CPU)"),
        ]
    )

    # -------------------------------------------------------------------------
    add_heading(doc, "12. How Everything Connects (Full Demo Flow)", 1)

    add_body(doc,
        "The following describes what happens when the system is used in a demo session.")
    steps = [
        ("Starting the app",
         "Running 'python main.py visualize --web' creates a solved CubieCube in memory "
         "and starts the Python HTTP server on port 8080. The browser loads index.html, "
         "builds 27 Three.js cubie objects in a solved configuration, and calls GET /state."),
        ("Scrambling",
         "The user clicks Scramble. The browser sends POST /scramble. The server calls "
         "generate_scramble() which produces a random WCA-legal sequence of 20 moves "
         "and applies them to the CubieCube, updating the four arrays. The new cp/co/ep/eo "
         "values are sent back as JSON. The browser calls syncState(), which computes the "
         "correct 3D quaternion for each piece using the 24-element rotation group."),
        ("Solving with Kociemba",
         "The user clicks Solve (Expert). The browser sends POST /solve. If a move history "
         "is available, the solution is instantly computed by reversing and inverting it. "
         "Otherwise bidirectional BFS is run. The solution list is returned as JSON, "
         "enqueued as animations, and played back one move at a time."),
        ("Solving with the AI",
         "The user clicks Solve (AI). The browser sends POST /solve_ai. The server calls "
         "encode_state() to convert the CubieCube to a 324-vector, feeds it through "
         "RubiksMLP, picks the highest-probability move, applies it, and repeats until "
         "solved or 50 steps have been attempted. The solution is animated the same way."),
        ("Benchmarking",
         "Running 'python main.py compare --max-depth 5' generates 50 random scrambles "
         "at each depth level, runs both solvers on each, and prints a table of solve "
         "rate, average moves, and average time side by side."),
    ]
    for step, desc in steps:
        p = doc.add_paragraph(style="List Number")
        run_bold = p.add_run(step + ": ")
        run_bold.bold = True
        p.add_run(desc)

    # -------------------------------------------------------------------------
    add_heading(doc, "Summary Table", 1)
    add_table(doc,
        ["Component", "File(s)", "Role"],
        [
            ("Cube model", "core/cube.py, core/cubie.py", "Mathematical state: cp, co, ep, eo arrays"),
            ("Move engine", "core/moves.py", "18 legal moves as permutation tables"),
            ("State encoder", "core/state_encoder.py", "Convert cube to 324-dim one-hot vector"),
            ("Scrambler", "utils/scramble.py", "Random WCA-style scramble generation"),
            ("Dataset generator", "data/dataset_generator.py", "Expert cloning + curriculum datasets"),
            ("Dataset stats", "data/dataset_stats.py", "Statistics on saved datasets"),
            ("Kociemba solver", "solvers/kociemba_solver.py", "Bidirectional BFS classical solver"),
            ("BFS solver", "solvers/bfs_solver.py", "Extended BFS with node statistics"),
            ("AI solver", "solvers/ai_solver.py", "MLP neural network solver"),
            ("Web server", "visualization/server.py", "REST API backend on port 8080"),
            ("Web frontend", "visualization/index.html", "Three.js interactive 3D viewer"),
            ("Desktop viewer", "visualization/controls.py, renderer3d.py", "Matplotlib 3D prototype"),
            ("Benchmarking", "experiments/ai_vs_kociemba.py", "AI vs Kociemba comparison"),
            ("Performance test", "experiments/kociemba_benchmark.py", "Kociemba standalone benchmark"),
            ("Tests", "tests/", "Automated validation suite (pytest)"),
            ("Entry point", "main.py", "CLI interface for all commands"),
        ]
    )

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    build()
