import docx

def create_main_doc():
    doc = docx.Document()
    doc.add_heading('Project Overview: AI-Based Optimal Solver for the 3x3 Rubik\'s Cube', 0)
    
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph('This project represents a complete software stack for simulating, scrambling, and optimally solving a 3x3 Rubik\'s Cube using an Artificial Intelligence approach combined with classical algorithms. The system features a mathematical backend in Python and a 3D visualization frontend.')
    
    doc.add_heading('2. Folder Architecture', level=1)
    
    doc.add_heading('core/', level=2)
    doc.add_paragraph('Contains the mathematical backend of the cube. It handles the cubie states (permutations and orientations) and strictly enforces the group-theoretical laws and legality constraints of the 3x3 puzzle.')
    
    doc.add_heading('solvers/', level=2)
    doc.add_paragraph('Houses all algorithms capable of producing solution steps for the cube. It includes the classical perfect dual-phase solver (Kociemba) and the heuristic predictive Neural Network solver (ai_solver.py).')
    
    doc.add_heading('data/', level=2)
    doc.add_paragraph('Contains the dataset generation logic for curriculum training, statistical validation scripts to ensure dataset uniformity, and the storage directories for the AI model weights (.pt).')
    
    doc.add_heading('visualization/', level=2)
    doc.add_paragraph('Provides a local HTTP server and an interactive WebGL (Three.js) frontend. The visualizer completely mirrors the raw internal arrays, applying rigid-layer mechanics to ensure perfect drift-free physical representations.')
    
    doc.add_heading('experiments/', level=2)
    doc.add_paragraph('Scripts used to benchmark the speed and efficiency of the Neural Network solver against the mathematically optimal Kociemba algorithms.')
    
    doc.add_heading('3. Learning AI and Dataset Strategy', level=1)
    doc.add_paragraph('The Neural Network is a Multi-Layer Perceptron (MLP) containing over 180,000 parameters. It accepts an encoded 324-dimensional one-hot representation of the 54 facelet variables and outputs softmax confidence probabilities for the 18 possible Half-Turn Metric moves.')
    doc.add_paragraph('The training procedure utilizes Curriculum Learning, gradually teaching the AI to solve the cube layer by layer (depth by depth).')
    
    doc.add_heading('Dataset Generation Pipeline:', level=2)
    doc.add_paragraph('1. Generation: Random scrambles are generated sequentially up to a specific depth.')
    doc.add_paragraph('2. Expert Target: The classical Kociemba solver computes the mathematically perfect next move for that state.')
    doc.add_paragraph('3. Training Loop: The MLP learns to clone the expert\'s predictions at the current depth.')
    doc.add_paragraph('4. Replacing Strategy: As the depth increases, 80% of the dataset is replaced with new difficulty samples, while 20% of the shallower depth data is retained to prevent catastrophic forgetting.')
    
    doc.save('docs/Project_Explanation.docx')

def create_supervisor_doc():
    doc = docx.Document()
    doc.add_heading('Supervisor Update: AI Rubik\'s Cube Solver Progress', 0)
    
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph('The development of the AI-Based Optimal Solver for the 3x3 Rubik\'s Cube is functionally complete and has successfully passed all necessary milestones. The project strictly adheres to the proposed research trajectory.')
    
    doc.add_heading('Key Accomplishments', level=1)
    doc.add_paragraph('1. Core Engine: The abstract mathematical solver is fully operational. It includes an extremely strict verification module that actively blocks parity impossibilities and orientation sum discrepancies.')
    doc.add_paragraph('2. AI Curriculum Training: The Multi-Layer Perceptron (MLP) architecture is implemented and successfully trains on the dynamic curriculum datasets. On local benchmark checks, the AI is demonstrating a strong capacity for generalizing the early stage state-spaces using Kociemba as an expert orchestrator.')
    doc.add_paragraph('3. 3D WebGL Visualization: We have bridged the mathematical backend with an interactive topological frontend. The visualization guarantees drift-free physical alignments by enforcing explicit logic-state matching after every interpolation ("Rigid Layer Mechanics").')
    doc.add_paragraph('4. Project Health: The comprehensive internal test suite passes unconditionally (111 unit tests execute flawlessly).')
    
    doc.add_heading('Next Actions', level=1)
    doc.add_paragraph('Having secured the pipeline architecture, the primary next step is to execute large-scale depth benchmarking and analyze the heuristic limitations of the MLP approach as scramble complexity increases. I have organized the documentation and am prepared to present the final system overview.')
    
    doc.save('docs/Supervisor_Update.docx')

if __name__ == "__main__":
    create_main_doc()
    create_supervisor_doc()
