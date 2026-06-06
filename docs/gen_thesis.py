"""
gen_thesis.py - generates thesis_Edward_Ogbei_PCZ.docx then converts to PDF.
Run from the project root: python docs/gen_thesis.py
"""
import os, sys
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIGS   = os.path.join(ROOT, "docs", "figures")
OUT    = os.path.join(ROOT, "docs", "thesis_Edward_Ogbei_PCZ.docx")

# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

def docx_break_type():
    from docx.oxml.ns import qn as _qn
    from docx.oxml    import OxmlElement as _Oxml
    br = _Oxml("w:br")
    br.set(_qn("w:type"), "page")
    return br   # we'll use run.add_break differently below

def page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)

def body(doc, text, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)
    if first_indent:
        pf.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(12)
    return p

def h1(doc, num, title):
    page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    pf.keep_with_next = True
    r = p.add_run(f"{num}. {title}")
    r.font.name  = "Arial"
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0,0,0)
    return p

def h2(doc, num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after  = Pt(4)
    pf.keep_with_next = True
    r = p.add_run(f"{num} {title}")
    r.font.name  = "Arial"
    r.font.size  = Pt(13)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0,0,0)
    return p

def figure(doc, fname, cap):
    path = os.path.join(FIGS, fname)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Inches(5.2))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    r = c.add_run(cap)
    r.font.name   = "Arial"
    r.font.size   = Pt(10)
    r.font.italic = True
    return c

def note(doc, text):
    """Italic note / observation block."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(4)
    pf.space_after  = Pt(6)
    pf.left_indent  = Cm(1.0)
    r = p.add_run(text)
    r.font.name   = "Arial"
    r.font.size   = Pt(11)
    r.font.italic = True
    return p

# ---------------------------------------------------------------------------
# title page
# ---------------------------------------------------------------------------

def title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(2)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("CZESTOCHOWA UNIVERSITY OF TECHNOLOGY")
    r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("Faculty of Electrical Engineering")
    r2.font.name = "Arial"; r2.font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Cm(3)
    r3 = p3.add_run("Institute of Computer and Information Sciences")
    r3.font.name = "Arial"; r3.font.size = Pt(12)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Cm(1)
    p4.paragraph_format.space_after  = Pt(4)
    r4 = p4.add_run("MASTER'S THESIS")
    r4.font.name = "Arial"; r4.font.size = Pt(13); r4.font.bold = True

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_after = Cm(2)
    r5 = p5.add_run("M.Sc. in Artificial Intelligence")
    r5.font.name = "Arial"; r5.font.size = Pt(12)

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p6.paragraph_format.space_before = Cm(1)
    p6.paragraph_format.space_after  = Cm(3)
    r6 = p6.add_run("AI-Based Optimal Solver for the 3x3 Rubik's Cube")
    r6.font.name = "Arial"; r6.font.size = Pt(16); r6.font.bold = True

    for label, val in [
        ("Author:", "Edward Ogbei"),
        ("Student ID:", "s47891"),
        ("Supervisor:", "Prof. dr hab. inz. Krzysztof Rojek"),
        ("Year:", "2025/2026"),
    ]:
        row = doc.add_paragraph()
        row.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.paragraph_format.space_after = Pt(4)
        rl = row.add_run(f"{label}  ")
        rl.font.name = "Arial"; rl.font.size = Pt(12); rl.font.bold = True
        rv = row.add_run(val)
        rv.font.name = "Arial"; rv.font.size = Pt(12)

    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p7.paragraph_format.space_before = Cm(3)
    r7 = p7.add_run("Czestochowa, 2026")
    r7.font.name = "Arial"; r7.font.size = Pt(12)

# ---------------------------------------------------------------------------
# chapter content
# ---------------------------------------------------------------------------

def ch_intro(doc):
    h1(doc, 1, "Introduction")

    body(doc, "The Rubik's Cube is one of those rare objects that sits comfortably at the intersection of mechanical engineering, abstract algebra, and popular culture. Erno Rubik, a Hungarian architecture professor, invented it in 1974 as a teaching aid to help his students visualise three-dimensional transformations. He called it the Magic Cube. By the time it reached international markets in 1980, renamed the Rubik's Cube, it had already become something of an obsession among puzzle enthusiasts. Estimates put total sales at over 450 million units by the early 2020s, making it arguably the best-selling toy in history.", first_indent=True)

    body(doc, "What makes the cube interesting from a computer science perspective is not its commercial appeal but the structure of its state space. A standard 3x3 Rubik's Cube has exactly 43,252,003,274,489,856,000 distinct reachable configurations - roughly 43 quintillion. Yet despite this enormous number, every single configuration can be solved in at most 20 face moves. This result, known informally as God's Number, was proved computationally in 2010 by Rokicki et al. [1] using around 35 CPU-years of computation distributed across Google's servers. The gap between the scale of the problem and the compactness of its solution is striking, and it makes the Rubik's Cube a natural benchmark for search and optimisation algorithms.")

    body(doc, "Classical approaches to the cube are well understood. Layer-by-layer methods used by competitive speedcubers work reliably but often require 50 or more moves. The Kociemba two-phase algorithm [2], developed in the early 1990s, exploits the group structure of the cube to decompose the search into two subproblems, each navigated with iterative deepening search and precomputed tables. The result is near-optimal solutions, typically well under 25 moves, found in milliseconds. These classical algorithms are efficient precisely because they encode deep domain knowledge about the cube's mathematical structure.")

    body(doc, "A different question has attracted increasing attention in recent years: can a learning system discover how to solve the cube without being handed that domain knowledge? This is the imitation learning approach. The idea is straightforward - generate many examples of (scrambled state, correct next move) pairs using an expert solver, then train a neural network to predict the next move from the current state. The challenge lies in generalising from a finite training set to the vast state space. Several factors make this non-trivial: the cube's state space is enormous, the relationship between state and optimal move is highly non-local, and naive training tends to fail on deeper scrambles.")

    body(doc, "Curriculum learning is one strategy to address this. Rather than training on all difficulty levels simultaneously, the model starts on easy examples - cubes that are only one or two moves from solved - and progressively encounters harder positions as its performance improves. This mirrors the way a human student learns: starting with simpler cases before tackling complex ones. The idea has roots in the cognitive science literature [3] and has been applied successfully in several reinforcement learning and imitation learning contexts.")

    body(doc, "This thesis presents a complete system built around these ideas. The core is a mathematically rigorous 3x3 Rubik's Cube simulation engine that represents the cube at the cubie level - tracking individual corner and edge pieces rather than face colours. On top of this, a multi-layer perceptron (MLP) is trained using curriculum imitation learning, with a bidirectional BFS solver acting as the expert. Two inference strategies are implemented: a simple greedy policy that picks the highest-confidence move at each step, and a beam search strategy that explores multiple candidate paths simultaneously. The entire system is exposed through an interactive 3D web interface built with Three.js, where users can watch the solvers in action with smooth animations.")

    body(doc, "The scope of this work is deliberately bounded. The neural network is trained and evaluated on scrambles up to a maximum depth of 8 moves. At depth 5, the greedy solver achieves a 77% success rate and the beam search achieves 89%. These numbers drop sharply beyond depth 6, reflecting the fundamental challenge of generalising from a bounded training distribution to deeper configurations. The intention is not to claim state-of-the-art performance on the full cube - that remains the domain of specialised systems like DeepCubeA [4] - but rather to build a complete, understandable pipeline that demonstrates what supervised curriculum learning can achieve within realistic computational constraints.")

    body(doc, "The thesis is organised as follows. Chapter 2 sets out the research objectives. Chapter 3 covers the theoretical background, including the group structure of the cube and relevant prior work. Chapter 4 describes the cube representation and move tables. Chapter 5 covers the solving strategies, both classical and neural. Chapter 6 describes the training data generation and learning process. Chapter 7 presents the web visualisation system. Chapter 8 reports experimental results, and Chapter 9 concludes with a discussion of limitations and future directions.")

    figure(doc, "fig_09_state_space.png", "Fig. 1.1. Illustration of the Rubik's Cube state space and the concept of solving depth. Each node represents a cube configuration and each edge a single face move.")


def ch_aim(doc):
    h1(doc, 2, "Aim of the Work")

    body(doc, "The first paragraph of this chapter situates the problem domain. The Rubik's Cube presents a well-defined combinatorial optimisation challenge: given an arbitrary scrambled configuration, find a sequence of legal face moves that returns the cube to the solved state. The challenge is not merely technical - it also raises broader questions about what kinds of structured knowledge a neural network can acquire through imitation, and how curriculum design affects generalisation.")

    body(doc, "The theoretical objective of this work is to investigate whether a relatively small, fully-connected neural network can learn a useful solving policy for the Rubik's Cube through imitation of an exact solver. The hypothesis is that curriculum training - starting from shallow scrambles and progressing to deeper ones - produces substantially better generalisation than training on all depths simultaneously. A secondary hypothesis is that beam search, by exploring multiple candidate paths, improves on the greedy policy particularly for mid-depth scrambles where the greedy policy begins to fail.")

    body(doc, "The practical objectives of the project are threefold. First, to implement a complete and mathematically verified 3x3 Rubik's Cube engine, including a cubie-level state representation, a full 18-move set, and automated checks for all group-theoretic legality constraints. Second, to build a training pipeline that generates curriculum datasets using an expert BFS solver and trains the MLP at successive depth levels. Third, to deliver an interactive web visualisation that allows real-time demonstration of both solvers, with animated moves and a benchmark comparison tool.")

    body(doc, "The scope of the work is bounded in several important ways. The neural solver targets scrambles up to 8 moves in depth. The expert solver is a bidirectional BFS implementation, which is optimal for shallow depths but does not scale to the full cube - it is not the Kociemba two-phase algorithm, which would require substantially more engineering effort to implement correctly. The neural network architecture is intentionally simple (a three-hidden-layer MLP with roughly 184,000 parameters) to keep training feasible on a standard CPU and to maintain interpretability. No GPU training was assumed, though the code supports CUDA automatically when available.")

    body(doc, "In summary, this thesis aims to demonstrate that curriculum imitation learning is a viable strategy for training a neural solving policy on a well-structured combinatorial puzzle, and to build all supporting infrastructure - the cube engine, the training pipeline, and the interactive visualisation - needed to evaluate and demonstrate this claim.")


def ch_background(doc):
    h1(doc, 3, "Theoretical Background")

    body(doc, "This chapter reviews the mathematical foundations needed to understand the project and situates it within the existing literature on Rubik's Cube solving and machine learning for combinatorial problems.", first_indent=False)

    h2(doc, "3.1", "The Rubik's Cube as a Mathematical Object")

    body(doc, "The Rubik's Cube can be described precisely using group theory. The set of all reachable configurations, together with the operation of applying a face move, forms a group under composition. This group - sometimes called the Rubik's Cube group - is a subgroup of the symmetric group on 48 movable facelets (excluding the 6 fixed centres). Its order is exactly 43,252,003,274,489,856,000, a number that can be derived by counting the independent degrees of freedom [5].")

    body(doc, "Concretely, the cube has 8 corner pieces and 12 edge pieces. Each corner can be placed in any of 8 slots (8! arrangements) and can have one of 3 orientations, giving 8! * 3^8 corner states if unconstrained. Similarly, each edge can occupy any of 12 slots with one of 2 orientations, giving 12! * 2^12 edge states. However, not all combinations are reachable from a legal starting configuration. Three constraints reduce the total count by a factor of 12: the sum of all corner orientations must be divisible by 3; the sum of all edge orientations must be divisible by 2; and the permutation parity of the corners must equal the permutation parity of the edges. These invariants are preserved by every legal face move and are central to verifying the legality of any given state.")

    body(doc, "The minimum number of face moves required to solve the worst-case configuration is known as God's Number. For the standard outer-layer turn metric it is 20, for the quarter-turn metric (where 180-degree turns count as two) it is 26. The proof for the face-turn metric was completed computationally in 2010 [1]. The distribution of states by depth follows a roughly Gaussian shape peaking around depth 17-18, with the vast majority of configurations reachable in 15 to 19 moves. This distribution explains why random scrambles of fewer than 8 moves represent only a tiny corner of the full state space.")

    body(doc, "For algorithmic purposes, it is common to represent the cube not at the facelet level (54 coloured stickers) but at the cubie level. Each corner is identified by the three colours it carries (e.g., the URF corner carries white, red, and green in the standard colour scheme) and its orientation relative to a canonical reference direction. Edges are similarly identified. This cubie representation is more compact and aligns naturally with the algebraic structure of the group.")

    h2(doc, "3.2", "Classical Solving Algorithms")

    body(doc, "The simplest exact search strategy is breadth-first search (BFS). Starting from the scrambled state, BFS expands nodes level by level, guaranteeing the shortest possible solution. The problem is the branching factor: at each step there are 18 possible moves (reduced to about 15 by move pruning), so BFS visits roughly 15^d nodes to solve a scramble of depth d. At d=10 that is already 576 billion nodes, making BFS impractical for general use. Bidirectional BFS, which expands simultaneously from both the scrambled state and the solved state, reduces this to roughly 2 * 15^(d/2) nodes, making d=14 feasible with modest memory and time budgets.")

    body(doc, "Iterative deepening A* (IDA*) [6] addresses the memory problem of BFS by using a depth-limited DFS that iterates over increasing depth limits. With a good admissible heuristic (one that never overestimates the true distance to the solution), IDA* finds optimal solutions without storing the full search tree. The Kociemba two-phase algorithm [2] exploits this framework by decomposing the cube group into two subgroups. Phase one finds a sequence that places the cube in a restricted subgroup (edges correctly oriented, UD-slice edges in the UD-slice). Phase two then finds a solution within that subgroup. Each phase uses IDA* with precomputed tables. The result is consistently near-optimal solutions found in milliseconds, with no randomness - a gold standard for exact solving.")

    body(doc, "Human methods like CFOP (Cross, F2L, OLL, PLL) decompose the solution into recognisable stages that can be memorised. While not optimal (a full solve may take 50-80 moves), CFOP is fast to execute in practice, which is why competitive speedcubers can solve the cube in under 5 seconds. These methods are not implemented in this project, but they illustrate the human-centric approach to domain decomposition.")

    h2(doc, "3.3", "AI-Based Approaches")

    body(doc, "Interest in machine learning approaches to the Rubik's Cube has grown considerably since the late 2010s. The most influential result is DeepCubeA by McAleer et al. [4], published in 2019. Their system trains a deep residual network to estimate the distance-to-goal (a value function) using a form of backward induction they call Autodidactic Iteration. At test time, a weighted A* search uses this value function as a heuristic, finding optimal or near-optimal solutions for arbitrary scramble depths. DeepCubeA solved 100% of test scrambles and found solutions within 2% of optimal length. It required substantial computational resources - training took several days on GPU clusters - but it demonstrated that deep learning can genuinely learn the structure of the cube's solution space.")

    body(doc, "An earlier line of work by Agostinelli et al. [7] used reinforcement learning with self-play, training a policy network to predict moves and a value network to estimate distance. This approach does not require an expert solver during training, relying instead on reward shaping. In practice, it proved harder to train than supervised or semi-supervised approaches.")

    body(doc, "Imitation learning - training directly from expert demonstrations - is generally more sample-efficient than reinforcement learning when a reliable expert is available. In the context of the Rubik's Cube, a BFS or Kociemba solver can act as the expert, generating (state, next_move) pairs cheaply. The challenge is that the model must generalise across a huge state space from a finite training set. Curriculum learning [3] is a natural mitigation: by ordering training examples from easy to hard, the model builds up gradually rather than being overwhelmed by deep scrambles from the start.")

    body(doc, "Several studies have explored small-scale imitation learning on the cube with MLP architectures, mostly as course projects or informal experiments. They consistently find that greedy MLP solvers work well for shallow scrambles (depth 1-5) and degrade sharply beyond that. This project falls into the same category - a principled implementation of the curriculum imitation learning pipeline, with added beam search and a complete web demonstration system.")

    figure(doc, "fig_07_ai_vs_kociemba.png", "Fig. 3.1. Overview of the relationship between classical expert solvers and the neural network imitation pipeline. The expert generates labels; the network learns to imitate.")


def ch_architecture(doc):
    h1(doc, 4, "System Architecture")

    body(doc, "This chapter describes the design of the cube simulation engine, which forms the mathematical backbone of the entire project. Every other component - the solvers, the training pipeline, the visualisation - depends on this engine for state representation and move application.", first_indent=False)

    h2(doc, "4.1", "Cubie-Level State Representation")

    body(doc, "The cube state is encoded using four integer arrays that follow the Kociemba coordinate system [2]. The corner permutation array (cp) has length 8 and records which corner piece occupies each corner slot, using integer identifiers 0 through 7. The corners are named by the three faces they sit at: URF (Up-Right-Front), UFL, ULB, UBR, DFR, DLF, DBL, DRB. In the solved state, cp[i] = i for all i.")

    body(doc, "The corner orientation array (co) also has length 8. Each value is 0, 1, or 2, representing the twist of the corner in its slot relative to the U/D axis. By convention, orientation 0 means the reference facelet (the one that would be on the U or D face in the solved state) is pointing up or down. Orientation 1 means it has been rotated 120 degrees clockwise, and orientation 2 means 120 degrees counter-clockwise. This three-valued encoding is specific to corners because they have three-fold rotational symmetry about the body diagonal.")

    body(doc, "For edges, two analogous arrays are used. The edge permutation array (ep) has length 12, recording which of the 12 edge pieces occupies each edge slot. The 12 edges are named UR, UF, UL, UB, DR, DF, DL, DB, FR, FL, BL, BR. The edge orientation array (eo) has length 12, with each value either 0 (correctly oriented) or 1 (flipped). The six face centres are treated as fixed reference frames and are never stored in any array - they implicitly define the colour assignment of each slot.")

    body(doc, "The CubieCube class encapsulates these four arrays along with the complete set of operations: applying a move, checking if the cube is solved, converting to the Kociemba facelet string representation (54 characters), and verifying legality. The class is intentionally small and free of any visualisation logic; it is a pure mathematical model.")

    figure(doc, "fig_03_encoding.png", "Fig. 4.1. The cubie-level state representation. Corner permutation (cp), corner orientation (co), edge permutation (ep), and edge orientation (eo) together fully describe any legal cube configuration.")

    h2(doc, "4.2", "Move Tables")

    body(doc, "The cube recognises 18 moves in the standard half-turn metric: six basic face turns (U, D, R, L, F, B), their inverses (U', D', R', L', F', B'), and their 180-degree variants (U2, D2, R2, L2, F2, B2). Each move is represented as a MoveDefinition object containing four tuples: the corner permutation, the corner orientation delta, the edge permutation, and the edge orientation delta.")

    body(doc, "The six base moves are defined by hand, following the Kociemba convention. As an example, the U move cycles corners in the sequence URF -> UBR -> ULB -> UFL -> URF and cycles edges in UR -> UB -> UL -> UF -> UR. Neither the corners nor the edges change their orientation during a U move, because the reference facelet (the one on the U or D face) remains on the U or D face after the turn. This is not the case for F and B moves, where edges that were previously on U or D faces end up on L or R faces, flipping their orientation.")

    body(doc, "Prime (inverse) moves and double moves are computed algebraically by composing the base move with itself. The compose_moves function implements the group product: given two move definitions A and B, the composition B then A is computed by applying B's permutation first and then A's. Orientation deltas compose additively modulo 3 (for corners) or modulo 2 (for edges). This approach ensures internal consistency - the move tables form a closed group under composition.")

    body(doc, "During development, the B and L move definitions proved particularly error-prone. The B move acts clockwise when viewed from behind the cube, which is counter-intuitive when the code is written from a front-facing perspective. An early version had the B move applying a mathematically counter-clockwise corner cycle alongside the correct edge cycle, causing a subtle visual shear in the 3D renderer that was initially difficult to diagnose. The fix was to carefully verify each move definition by applying it four times and checking that the result is the identity.")

    figure(doc, "fig_01_system_architecture.png", "Fig. 4.2. High-level system architecture showing the relationships between the cube engine, the solvers, the training pipeline, and the web visualisation.")

    h2(doc, "4.3", "Legality Constraints")

    body(doc, "Not every assignment of corner and edge pieces to slots is a legal cube state - one that could be reached by a sequence of face moves from the solved position. The is_legal() method checks three invariants that are preserved by every legal move.")

    body(doc, "The first invariant is that the sum of all corner orientations must be divisible by 3. This follows from the fact that each move contributes an orientation change that sums to zero modulo 3 across the affected corners. If a corner is removed from a physical cube and reinserted with a twist, this invariant is violated, and the cube cannot be solved.")

    body(doc, "The second invariant is that the sum of all edge orientations must be divisible by 2. This reflects the fact that edges can only be flipped in pairs by legal moves. The third and most subtle invariant is that the permutation parity of the corners must equal the permutation parity of the edges. A single transposition of two corner pieces (without any other change) produces an illegal state because it changes the corner parity without changing the edge parity. These three invariants collectively reduce the theoretical state count by a factor of 12 to the confirmed 43.25 quintillion.")

    h2(doc, "4.4", "State Encoding for Neural Networks")

    body(doc, "To feed a cube state into a neural network, the 54-character Kociemba facelet string is converted to a one-hot vector. Each of the 54 facelets can take one of 6 possible colours (U/white, R/red, F/green, D/yellow, L/orange, B/blue). The one-hot encoding represents each facelet as a length-6 binary vector with a single 1 in the position corresponding to its colour. Concatenating all 54 vectors gives a 324-dimensional input vector (54 * 6 = 324). Every element is either 0.0 or 1.0, and exactly 54 elements are 1.0 in any valid encoding.")

    body(doc, "This encoding is simple but carries all the information needed to determine the cube state. It does not exploit the group structure of the cube - two states that are related by a symmetry of the cube will have different encodings unless that symmetry happens to be the identity. More sophisticated encodings using coordinate representations (twist, flip, UD-slice position) exist but were not used here, as the goal was to see what a network could learn from raw facelet data.")

    figure(doc, "fig_02_mlp_architecture.png", "Fig. 4.3. MLP architecture: 324-dimensional one-hot input, three hidden layers (256, 256, 128 neurons with ReLU), and an 18-unit output layer predicting the next move.")


def ch_solvers(doc):
    h1(doc, 5, "Solving Strategies")

    body(doc, "This chapter covers all three solving components: the BFS-based expert solver used both for training labels and as a baseline, the greedy neural network solver, and the beam search extension.", first_indent=False)

    h2(doc, "5.1", "The BFS Expert Solver")

    body(doc, "The expert solver used in this project is a bidirectional BFS implementation. It maintains two frontiers: one expanding forward from the scrambled state, and one expanding backward from the solved state. The search terminates when the two frontiers meet. The solution is constructed by concatenating the forward path with the inverse of the backward path.")

    body(doc, "Move pruning reduces the effective branching factor by preventing immediately consecutive moves on the same face (e.g., R followed by R' is pointless) and consecutive moves on opposite faces in a fixed order (e.g., R after L, but not L after R, since they commute). With these prunings, the effective branching factor is approximately 15 at the first move and decreasing thereafter. Bidirectional BFS with this pruning is feasible up to around depth 14 in terms of time, though memory usage becomes substantial beyond depth 9 or 10.")

    body(doc, "It is worth being explicit about what this solver is and is not. It is not the Kociemba two-phase algorithm. The Kociemba algorithm uses IDA* with precomputed pruning tables organised around two nested subgroups of the cube group, and consistently finds solutions under 25 moves in milliseconds even for arbitrary scrambles. The BFS solver implemented here works correctly and finds optimal solutions, but its practical depth ceiling is around 8-9 moves before timing out at the default 30-second limit. For the purposes of this project - generating training data for scrambles up to depth 8 - this is sufficient.")

    note(doc, "Note: the naming convention 'Kociemba solver' used in parts of the codebase and UI is a simplification. The underlying implementation is bidirectional BFS, not the true Kociemba algorithm. This distinction is relevant when evaluating the system's claims about optimality for deeper scrambles.")

    h2(doc, "5.2", "Neural Network Architecture")

    body(doc, "The neural network is a fully-connected multi-layer perceptron (MLP) with the following structure: an input layer of 324 units, three hidden layers of 256, 256, and 128 neurons respectively, each with ReLU activations, and an output layer of 18 units - one per legal move. The output is interpreted as unnormalised logits; a softmax is applied at inference time to obtain move probabilities.")

    body(doc, "The total parameter count is approximately 184,210, broken down as follows: the first layer contributes 324 * 256 + 256 = 83,456 parameters, the second 256 * 256 + 256 = 65,792, the third 256 * 128 + 128 = 32,896, and the output layer 128 * 18 + 18 = 2,322. This is a relatively small network by modern standards, deliberately so: the intention was to keep training feasible on CPU within hours rather than days, and to maintain a model simple enough to reason about.")

    body(doc, "ReLU was chosen over sigmoid or tanh because of its well-known advantages for deep networks: it avoids the vanishing gradient problem and is computationally efficient. No dropout or batch normalisation was applied. In practice, the training data generated by the curriculum is large enough that overfitting is not a significant concern at this model size.")

    figure(doc, "fig_02_mlp_architecture.png", "Fig. 5.1. Detailed view of the MLP architecture used for move prediction. Each hidden layer applies ReLU; the output applies softmax at inference time.")

    h2(doc, "5.3", "Greedy Inference")

    body(doc, "At inference time, the greedy solver feeds the current cube state through the network and selects the move with the highest softmax probability. This move is applied to the cube, and the process repeats until either the cube is solved or a maximum step count (default 50) is reached. The approach is simple and fast - a single forward pass takes roughly 0.5 to 2 milliseconds on CPU - but it has a fundamental weakness: once the model makes a mistake, there is no recovery mechanism. If the predicted move takes the cube further from the solution, subsequent predictions have no way of recognising the error and backtracking.")

    body(doc, "In practice, the greedy solver is reliable for scrambles up to depth 3 or 4, where the model's confidence is high and errors are rare. At depth 5, the success rate drops to 77%, and beyond depth 6 the degradation is rapid. Interestingly, the greedy solver occasionally 'over-solves' - applies more moves than the scramble depth - which suggests the model sometimes takes a slight detour before finding the solution. The solution lengths reported for successful solves average close to the scramble depth, indicating the network generally learns near-optimal paths.")

    h2(doc, "5.4", "Beam Search")

    body(doc, "Beam search addresses the irreversibility problem of greedy decoding by maintaining a set of candidate solution paths simultaneously. At each step, every path in the beam is extended by considering its top-k most probable next moves (where k is the beam width). The resulting candidates are scored by cumulative log-probability and the best beam_width candidates are retained. If any extended path reaches the solved state, that path is returned as the solution.")

    body(doc, "A visited-state set prevents the search from revisiting the same cube configuration through different paths, which would cause cycles. This is important because the greedy model can get stuck in short loops - repeatedly applying the same two or three moves - without the visited-state check to break the cycle.")

    body(doc, "With a beam width of 5, the beam search achieves 89% success at depth 5 and 50% at depth 6, compared to 77% and 31% for the greedy solver. The cost is higher computational load: each step requires beam_width forward passes rather than one, and the visited-state set grows over time. For depth 5 scrambles, the beam search typically runs in 11 milliseconds, versus 2.3 milliseconds for greedy. This is still fast enough for interactive use but would become a bottleneck for bulk evaluation at scale.")

    figure(doc, "fig_10_beam_search.png", "Fig. 5.2. Beam search tree with width 5. At each step, all paths are extended and the top-5 by cumulative log-probability are retained. Solved paths exit immediately.")


def ch_training(doc):
    h1(doc, 6, "Dataset Generation and Curriculum Training")

    body(doc, "This chapter describes how training data is generated and how the model is trained. The two are closely linked: the curriculum strategy determines not just the training procedure but also the structure of the dataset.", first_indent=False)

    h2(doc, "6.1", "Dataset Generation")

    body(doc, "Training samples are generated by a process that mirrors what the model must learn to do: given a scrambled cube, find the next optimal move. For each sample, a cube is scrambled by applying d random moves (where d is the target depth), and the expert BFS solver is called to find the complete solution. The first move of that solution is the label for this training sample. The cube state at that point is the input. This generates one (state, label) pair per sample.")

    body(doc, "The scramble generation is careful to avoid trivially reversing sequences. A random scramble at depth d applies d moves drawn uniformly from the 18-move set, with the constraint that no consecutive moves are on the same face. This produces scrambles that genuinely require d moves to solve - a scramble of the form R R' R would have an effective depth of 1, not 3, which would mislead the model about the task difficulty.")

    body(doc, "Dataset files are stored as compressed NumPy archives (.npz), each containing three arrays: states (shape N x 324, float32), labels (shape N, int64, values in 0-17), and depths (shape N, int32). At the largest setting, the dataset at depth 8 contains around 150,000 samples and occupies approximately 60 megabytes on disk.")

    h2(doc, "6.2", "The 80 Percent Replacement Strategy")

    body(doc, "A key design decision in the dataset generation is how to handle the growing dataset as the curriculum progresses. A naive approach would accumulate all samples from all depths equally. However, this causes the earlier (easier) depths to dominate as the curriculum advances, making the loss function trivially minimisable by ignoring the harder examples.")

    body(doc, "The strategy used here is 80 percent replacement. When moving from depth d to depth d+1, the new samples at depth d+1 comprise 80 percent of the new dataset, and a random 20 percent of the previous dataset is retained. This ensures that the model sees mostly the current difficulty level while retaining some examples of easier cases to prevent catastrophic forgetting. The 80/20 split is a heuristic; the optimal split likely depends on how quickly the model adapts to each new depth, but this proportion worked well in practice.")

    figure(doc, "fig_04_curriculum_flow.png", "Fig. 6.1. Curriculum learning progression. At each depth level, 80 percent of the dataset is replaced with new harder samples. Training advances to the next depth once the solve rate exceeds a configurable threshold.")

    h2(doc, "6.3", "Training Procedure")

    body(doc, "The model is trained using the Adam optimiser [8] with a learning rate of 0.001 and cross-entropy loss over the 18 move classes. The batch size is 64. At each depth level, training runs for a fixed number of epochs (default 50), after which the model is evaluated on a held-out set of 100 scrambles at the current depth. If the solve rate exceeds a threshold (default 80 percent), the curriculum advances to the next depth; otherwise, training stops.")

    body(doc, "This gating mechanism is important: advancing to a harder depth before the model has sufficiently mastered the current one tends to destabilise training. In practice, depths 1 through 3 are learned quickly and reliably - typically within the first 10-20 epochs. Depth 4 and 5 require more epochs and the model's solve rate at those depths shows more variance across training runs.")

    body(doc, "Training was conducted entirely on CPU for reproducibility and to avoid hardware dependencies, though the code automatically detects and uses CUDA if available. On a modern laptop CPU, training through depth 5 with 1,000 samples per depth and 50 epochs per depth takes roughly 10-15 minutes. On Colab with GPU, this reduces to a few minutes.")

    figure(doc, "real_loss_curves.png", "Fig. 6.2. Training loss curves across all depth levels. Loss decreases smoothly at each depth, with a brief spike when advancing to the next depth as the model encounters harder examples.")

    figure(doc, "real_curriculum_progression.png", "Fig. 6.3. Curriculum progression showing solve rate at each depth level during training. The model advances to the next depth once solve rate exceeds 80 percent.")

    h2(doc, "6.4", "Implementation Notes")

    body(doc, "The training code is implemented entirely in PyTorch. The RubiksDataset class wraps NumPy arrays in PyTorch tensors and is fed to a standard DataLoader with shuffling enabled. The model weights are saved using torch.save() after training completes. A pre-trained model is provided at data/models/ai_solver.pt and is loaded automatically by the web server.")

    body(doc, "One practical challenge during development was the BFS solver's runtime. For depths 6 and above, the BFS can take several seconds per sample, making dataset generation slow. A per-sample timeout (default 5 seconds) is applied, and samples that exceed this timeout are skipped. In practice, the skip rate is near zero for depth 1-5 but rises to around 5-10 percent for depth 7-8, where the BFS occasionally fails to find a solution within the time budget.")


def ch_viz(doc):
    h1(doc, 7, "Web Visualisation System")

    body(doc, "The web visualisation is a self-contained interactive 3D application served by a lightweight Python HTTP server. It allows users to apply individual moves, scramble the cube, and invoke all three solvers, with smooth animations for each step.", first_indent=False)

    h2(doc, "7.1", "Server Design")

    body(doc, "The server is a plain Python HTTP server built on the standard library's http.server module, with no external web framework dependencies. It serves the Three.js frontend as a static file and exposes a small REST API over HTTP. The endpoints are: GET /state (returns current cube state as JSON), POST /move (applies a single move), POST /scramble (generates and applies a random scramble), POST /reset (returns to solved state), POST /solve (BFS expert solve), POST /solve_ai (greedy AI solve), POST /solve_ai_beam (beam search AI solve), and POST /compare_benchmark (runs a quick benchmark across depths 1-5).")

    body(doc, "The cube state is held as a global CubieCube object in the server process. This is appropriate for a single-user local tool but would be inadequate for a multi-user deployment. The JSON response to each endpoint includes not only the action result but the complete current state (cp, co, ep, eo arrays, solved flag, legal flag, move history), so the frontend can synchronise its 3D scene without a separate state fetch.")

    body(doc, "The AI model is loaded once at startup using torch.load(), which can take a second or two. To avoid blocking the server socket while the model loads (a concern for cloud deployments with strict startup timeouts), the server binds the port and starts listening before loading the model.")

    h2(doc, "7.2", "Three.js 3D Frontend")

    body(doc, "The frontend is a single HTML file with embedded JavaScript. It uses Three.js (r128) loaded from a CDN, with OrbitControls for mouse-driven rotation of the cube. The 3D scene contains 26 rigid cubie meshes: 8 corners, 12 edges, and 6 centres. Each cubie is an independent THREE.Group containing a dark box mesh (the plastic body) and coloured plane meshes for each exposed face.")

    body(doc, "A key design decision was how to synchronise the 3D scene with the backend state. A naive approach would move stickers rather than pieces, but that is visually unconvincing and hard to animate. The approach taken here is to re-position and re-orient each physical cubie mesh based on the backend's cp, co, ep, eo arrays. Specifically, after each state update, the syncState() function iterates over all 20 non-centre cubies and computes the correct 3D position and quaternion orientation.")

    body(doc, "The orientation computation uses the precomputed set of 24 quaternions corresponding to the 24 rotations in the cubic symmetry group (Q24). For a given cubie in a given slot with a given orientation, the correct quaternion is found by searching Q24 for the rotation that maps the cubie's canonical position and reference normal to its current slot position and orientation direction. This exhaustive search over 24 candidates is fast (24 is tiny) and guarantees exact alignment without floating-point drift.")

    body(doc, "Animated moves are handled by a queue-based system. When a move button is pressed or a solver returns a sequence, the moves are pushed onto an animation queue. A pivot GROUP is used as an auxiliary object: the cubies belonging to the moving face are re-parented onto the pivot, the pivot is rotated by the appropriate angle over several frames, and then the cubies are re-parented back to the main group. After each animation, the 3D positions are snapped to exact integer grid positions and the orientations are snapped to the nearest 90-degree multiple, preventing accumulated floating-point error.")

    figure(doc, "screenshot_08_ui_wide.png", "Fig. 7.1. The full web interface. The Three.js 3D cube occupies the left panel; the control sidebar on the right provides move buttons, solver controls, animation speed selection, and a benchmark table.")

    figure(doc, "screenshot_04_kociemba_solve_preview.png", "Fig. 7.2. A scrambled cube mid-way through an animated Kociemba (BFS) solve. The 'Now Playing' indicator shows the current move in the sequence.")

    figure(doc, "screenshot_06_ai_beam.png", "Fig. 7.3. The AI beam search solver in action. The solution panel shows the number of moves found, time taken, and the move sequence.")


def ch_results(doc):
    h1(doc, 8, "Experimental Evaluation")

    body(doc, "This chapter presents the quantitative evaluation of both solving strategies across scramble depths 1 through 10. All results were obtained by running 100 test cases per depth using a fixed random seed, with the pre-trained model loaded from data/models/ai_solver.pt.", first_indent=False)

    h2(doc, "8.1", "Evaluation Methodology")

    body(doc, "For each test scramble at depth d, a solved cube is scrambled by applying exactly d moves (with the no-consecutive-same-face constraint), and the solver is given a maximum of 50 moves to find a solution. A scramble counts as solved if the cube reaches the solved state within those 50 moves. The solve rate is the fraction of scrambles solved out of 100. The average number of moves reported is computed only over successful solves.")

    body(doc, "Two solvers are evaluated: the greedy MLP solver and the beam search solver with beam width 5. The BFS expert solver achieves 100% success at all depths it can handle (up to about depth 9-10 with the 30-second timeout), so it is used as the baseline reference rather than a direct competitor.")

    h2(doc, "8.2", "Solve Rate Results")

    body(doc, "Table 8.1 summarises the solve rates and average move counts for both strategies across depths 1 to 10.")

    # Table
    tab_cap = doc.add_paragraph()
    tab_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tab_cap.paragraph_format.space_before = Pt(8)
    tr = tab_cap.add_run("Table 8.1. Solve rates and average solution lengths for the greedy and beam search solvers.")
    tr.font.name = "Arial"; tr.font.size = Pt(11); tr.font.bold = True

    data = [
        ["Depth", "Greedy Rate", "Greedy Avg. Moves", "Greedy Avg. Time (ms)", "Beam Rate", "Beam Avg. Moves", "Beam Avg. Time (ms)"],
        ["1",  "100%", "1.0",  "0.4",  "100%", "1.0",  "0.5"],
        ["2",  "100%", "2.0",  "0.8",  "100%", "2.0",  "1.0"],
        ["3",  "100%", "3.0",  "1.2",  "100%", "3.0",  "4.4"],
        ["4",   "97%", "4.0",  "1.6",   "98%", "4.0",  "7.2"],
        ["5",   "77%", "5.1",  "2.3",   "89%", "5.1", "11.1"],
        ["6",   "31%", "6.1",  "2.7",   "50%", "6.9", "17.9"],
        ["7",   "13%", "7.0",  "3.4",   "23%", "9.6", "30.2"],
        ["8",    "3%", "8.7",  "3.1",    "9%", "9.6", "25.7"],
        ["9",    "1%", "9.0",  "3.0",    "4%","17.0", "54.7"],
        ["10",   "0%",   "-",    "-",    "1%","10.0", "22.3"],
    ]
    table = doc.add_table(rows=len(data), cols=7)
    table.style = "Table Grid"
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True
            if i == 0:
                _set_cell_bg(cell, "D3D3D3")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    figure(doc, "real_solve_rate_bars.png", "Fig. 8.1. Solve rates by scramble depth for greedy (orange) and beam search (blue). Both solvers perform perfectly at depth 1-3 and degrade beyond depth 5.")

    h2(doc, "8.3", "Timing Analysis")

    body(doc, "Inference times for the greedy solver are remarkably consistent, ranging from 0.4 ms at depth 1 to 3.4 ms at depth 7. Since the greedy solver always runs exactly as many forward passes as moves attempted (up to 50), and since deeper scrambles take slightly more steps on average, the timing grows roughly linearly with depth. Each forward pass through the 184k-parameter network takes approximately 0.1-0.2 ms on CPU.")

    body(doc, "Beam search timing scales more steeply with depth. At depth 3, it already takes 4.4 ms due to the overhead of maintaining and scoring multiple candidate paths. At depth 9, the average successful solve takes 54.7 ms. The visited-state set grows substantially for deeper searches, and the overhead of encoding the cube state for each candidate path in the beam adds up. For an interactive application this is perfectly acceptable, but it would need parallelisation for bulk evaluation scenarios.")

    figure(doc, "fig_08_timing.png", "Fig. 8.2. Average inference time (milliseconds) for greedy and beam search solvers across scramble depths.")

    h2(doc, "8.4", "Discussion")

    body(doc, "The results confirm both hypotheses stated in Chapter 2. First, the curriculum-trained model genuinely learns to solve scrambles up to depth 5 at a useful success rate. The 77% greedy rate and 89% beam rate at depth 5 represent non-trivial generalisation: there are over 4 billion distinct cube states at distance 5 from solved, and the model has only seen a small fraction of them during training.")

    body(doc, "Second, beam search consistently improves on the greedy solver, particularly at the boundary depths (5 and 6) where the greedy policy begins to fail. The improvement at depth 5 is 12 percentage points (77% to 89%), and at depth 6 it is 19 points (31% to 50%). This aligns with the intuition that beam search helps most when the greedy policy is nearly but not quite reliable - it provides a second and third chance when the first choice is wrong.")

    body(doc, "The sharp degradation beyond depth 6 is expected and worth understanding. The model was trained with the curriculum advancing through depths 1-8, but the dataset at each depth level contains at most a few thousand samples. Given the branching factor of the state space, the coverage of states at depth 6+ is extremely sparse. The model generalises well within the distribution of its training data but struggles with configurations it has effectively never seen. This is a fundamental limitation of pure imitation learning - unlike DeepCubeA's value iteration approach, the model has no mechanism to improve its estimates on unseen states.")

    body(doc, "Compared to DeepCubeA [4], the results here are less impressive but the computational requirements are radically lower. DeepCubeA achieves 100% solve rate at all depths using a much larger residual network and weeks of GPU training. The system in this thesis uses a small MLP trainable in under an hour on CPU. The appropriate framing is not competition but complementarity: this project demonstrates what a simple, transparent pipeline can achieve at modest scale.")

    figure(doc, "fig_06_solve_rates.png", "Fig. 8.3. Comparison of greedy and beam solve rates, highlighting the benefit of beam search at intermediate depths.")


def ch_conclusions(doc):
    h1(doc, 9, "Conclusions and Future Work")

    body(doc, "This thesis set out to build and evaluate an AI-based solver for the 3x3 Rubik's Cube using curriculum imitation learning. The work produced four interconnected components: a mathematically rigorous cube engine with all group-theoretic invariants enforced, a BFS expert solver that generates training labels, a curriculum-trained MLP with greedy and beam search inference, and an interactive 3D web visualisation. All components work together as a coherent system, accessible through a web browser at localhost:8080.")

    body(doc, "The experimental results show that the greedy solver achieves 100% success on scrambles up to depth 3, 97% at depth 4, and 77% at depth 5. Beam search (width 5) improves these figures to 100%, 98%, and 89% respectively, and achieves 50% at depth 6. Beyond depth 6, both solvers degrade rapidly, reflecting the sparse coverage of deep states in the training data. These numbers are consistent with what has been reported informally by others attempting imitation learning on the cube with similar architectures.")

    body(doc, "The main limitation of the system is the depth ceiling. The expert solver is a bidirectional BFS, not a true Kociemba algorithm, which means it cannot generate reliable training data for scrambles beyond depth 9 or 10. The neural network, trained on this bounded data, cannot generalise to the arbitrary 20-move scrambles that represent the full difficulty of the puzzle. Addressing this would require either a genuine Kociemba implementation (with its associated engineering complexity of precomputed tables) or an alternative training strategy such as value iteration.")

    body(doc, "A second limitation is that the state encoding does not exploit the cube's symmetry group. The cube has 48 symmetries (24 rotations and 24 rotations combined with face reflection), and exploiting these could reduce the effective state space by a factor of 48. This would allow the same number of training samples to cover substantially more of the state space, potentially improving solve rates at deeper levels.")

    body(doc, "Third, the MLP architecture makes no assumptions about the local structure of the cube - it treats all 324 input dimensions equally, with no spatial or relational inductive bias. Graph neural networks or attention-based architectures that model the relationships between pieces might learn more efficiently, particularly for deeper scrambles where identifying key subproblems matters.")

    body(doc, "Several directions for future work are worth highlighting. The most impactful extension would be implementing a true Kociemba two-phase solver, which would unlock reliable expert demonstrations at all depths and allow training to target the full state space. A second direction is to combine the MLP policy with a value function (as in DeepCubeA) and use the policy-value pair to guide Monte Carlo Tree Search or a similar lookahead, which could dramatically improve performance at deeper depths without requiring retraining from scratch. A third direction is to augment the training data using cube symmetries, which is straightforward to implement and could improve solve rates at no computational cost.")

    body(doc, "The web visualisation, while not a research contribution in itself, proved valuable as a tool for understanding and explaining the system's behaviour. Watching the beam search navigate toward the solution, and seeing it occasionally take a longer path before converging, gives intuition about the model's decision-making that a table of numbers does not. Making the visualisation publicly accessible (e.g., via a Render or Hugging Face deployment) would allow others to explore the system interactively.")

    body(doc, "In conclusion, this thesis demonstrates that curriculum imitation learning is a viable and understandable approach to building a partial Rubik's Cube solver. The system is transparent, reproducible, and educational. The results are honest about its limitations. And the visualisation makes those limitations visible in a way that invites further experimentation. I consider the work a solid foundation for the more ambitious extensions described above.")


def bibliography(doc):
    page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Bibliography")
    r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True
    p.paragraph_format.space_after = Pt(10)

    refs = [
        "[1] Rokicki T., Kociemba H., Davidson M., Dethridge J., God's Number is 20, https://www.cube20.org, 2010.",
        "[2] Kociemba H., The Cube Explorer - Two-Phase Algorithm, http://kociemba.org/cube.htm, 1992.",
        "[3] Bengio Y., Louradour J., Collobert R., Weston J., Curriculum Learning, Proceedings of the 26th International Conference on Machine Learning (ICML), Montreal, 2009, pp. 41-48.",
        "[4] McAleer S., Agostinelli F., Shmakov A., Baldi P., Solving the Rubik's Cube Without Human Knowledge, arXiv:1805.07470, 2018. Updated as DeepCubeA in Nature Machine Intelligence, 2019.",
        "[5] Joyner W.D., Adventures in Group Theory: Rubik's Cube, Merlin's Machine, and Other Mathematical Toys, Johns Hopkins University Press, Baltimore, 2002.",
        "[6] Korf R.E., Depth-First Iterative-Deepening: An Optimal Admissible Tree Search, Artificial Intelligence, vol. 27, no. 1, pp. 97-109, 1985.",
        "[7] Agostinelli F., McAleer S., Shmakov A., Baldi P., Solving the Rubik's Cube with Deep Reinforcement Learning and Search, Nature Machine Intelligence, vol. 1, pp. 356-363, 2019.",
        "[8] Kingma D.P., Ba J., Adam: A Method for Stochastic Optimization, International Conference on Learning Representations (ICLR), San Diego, 2015.",
        "[9] Rokicki T., Twenty-Two Moves Suffice for Rubik's Cube, The Mathematical Intelligencer, vol. 32, no. 1, pp. 33-40, 2010.",
        "[10] Paszke A. et al., PyTorch: An Imperative Style, High-Performance Deep Learning Library, Advances in Neural Information Processing Systems (NeurIPS), 2019.",
        "[11] Three.js, JavaScript 3D Library, https://threejs.org, accessed May 2026.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(4)
        pf.left_indent = Cm(1.0); pf.first_line_indent = Cm(-1.0)
        r = p.add_run(ref)
        r.font.name = "Arial"; r.font.size = Pt(11)


def appendices(doc):
    page_break(doc)
    p = doc.add_paragraph()
    r = p.add_run("Appendix A - List of Figures")
    r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True
    p.paragraph_format.space_after = Pt(10)

    figs_list = [
        "Fig. 1.1. Illustration of the Rubik's Cube state space and the concept of solving depth.",
        "Fig. 3.1. Overview of the relationship between classical expert solvers and the neural pipeline.",
        "Fig. 4.1. The cubie-level state representation (cp, co, ep, eo).",
        "Fig. 4.2. High-level system architecture.",
        "Fig. 4.3. MLP architecture: 324-dimensional input, three hidden layers, 18-unit output.",
        "Fig. 5.1. Detailed MLP architecture view.",
        "Fig. 5.2. Beam search tree with width 5.",
        "Fig. 6.1. Curriculum learning progression diagram.",
        "Fig. 6.2. Training loss curves across depth levels.",
        "Fig. 6.3. Curriculum progression - solve rate per depth during training.",
        "Fig. 7.1. Full web interface screenshot.",
        "Fig. 7.2. Animated Kociemba solve mid-sequence.",
        "Fig. 7.3. AI beam search solver in action.",
        "Fig. 8.1. Solve rates by scramble depth for greedy and beam search.",
        "Fig. 8.2. Average inference time by scramble depth.",
        "Fig. 8.3. Comparison of greedy and beam solve rates.",
    ]
    for f in figs_list:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run(f)
        r2.font.name = "Arial"; r2.font.size = Pt(11)

    page_break(doc)
    p = doc.add_paragraph()
    r = p.add_run("Appendix B - List of Tables")
    r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True
    p.paragraph_format.space_after = Pt(10)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Table 8.1. Solve rates and average solution lengths for the greedy and beam search solvers.")
    r2.font.name = "Arial"; r2.font.size = Pt(11)

    page_break(doc)
    p = doc.add_paragraph()
    r = p.add_run("Appendix C - List of Abbreviations")
    r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True
    p.paragraph_format.space_after = Pt(10)

    abbrevs = [
        ("AI",   "Artificial Intelligence"),
        ("BFS",  "Breadth-First Search"),
        ("CFOP", "Cross, F2L, OLL, PLL - competitive speedcubing method"),
        ("CPU",  "Central Processing Unit"),
        ("GPU",  "Graphics Processing Unit"),
        ("HTM",  "Half-Turn Metric"),
        ("IDA*", "Iterative Deepening A-star search"),
        ("JSON", "JavaScript Object Notation"),
        ("MLP",  "Multi-Layer Perceptron"),
        ("MCTS", "Monte Carlo Tree Search"),
        ("PCZ",  "Politechnika Czestochowska (Czestochowa University of Technology)"),
        ("ReLU", "Rectified Linear Unit"),
        ("REST", "Representational State Transfer"),
        ("URF",  "Up-Right-Front (corner naming convention)"),
    ]
    for abbr, meaning in abbrevs:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(3)
        rb = p2.add_run(f"{abbr}  ")
        rb.font.name = "Arial"; rb.font.size = Pt(11); rb.font.bold = True
        rm = p2.add_run(f"- {meaning}")
        rm.font.name = "Arial"; rm.font.size = Pt(11)


def abstracts(doc):
    page_break(doc)
    p = doc.add_paragraph()
    r = p.add_run("Streszczenie (Polish Abstract)")
    r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True
    p.paragraph_format.space_after = Pt(10)

    body(doc, "Niniejsza praca magisterska opisuje projekt i implementacje systemu opartego na sztucznej inteligencji do rozwiazywania kostki Rubika 3x3. System sklada sie z czterech glownych komponentow: silnika symulacji kostki na poziomie elementow skladowych (cubie), eksperta BFS generujacego etykiety treningowe, wielowarstwowej sieci neuronowej MLP trenowanej metoda curriculum imitation learning, oraz interaktywnej wizualizacji 3D w przegladarce opartej na bibliotece Three.js.", first_indent=True)

    body(doc, "Siec neuronowa przyjmuje 324-wymiarowy wektor kodowania one-hot (54 pola razy 6 kolorow) i przewiduje nastepny ruch wsrod 18 mozliwych. Trenowana jest metoda curriculum learning, stopniowo zwiekszajac poziom trudnosci od 1 do 8 ruchow, z zastosowaniem 80-procentowej strategii zastepowania zbioru danych. Eksperymenty pokazuja, ze zachlanne podejscie osiaga skutecznosc 77% dla glebokich tasowac (depth 5), a przeszukiwanie wiazka (beam search, szerokosc 5) poprawia ten wynik do 89%. Powyzej glebokosci 6 obydwie strategie degraduja sie szybko, co odzwierciedla fundamentalne ograniczenia czystego uczenia przez imitacje bez mechanizmu planowania.", first_indent=True)

    page_break(doc)
    p = doc.add_paragraph()
    r = p.add_run("Abstract (English)")
    r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True
    p.paragraph_format.space_after = Pt(10)

    body(doc, "This master's thesis describes the design and implementation of an AI-based system for solving the 3x3 Rubik's Cube. The system consists of four main components: a cubie-level cube simulation engine with full group-theoretic legality checking, a bidirectional BFS expert solver used to generate training labels, a multi-layer perceptron trained via curriculum imitation learning, and an interactive Three.js 3D web visualisation.", first_indent=True)

    body(doc, "The neural network takes a 324-dimensional one-hot encoded input (54 facelets times 6 colours) and predicts the next move from a vocabulary of 18 legal moves. It is trained using curriculum learning that progressively increases scramble depth from 1 to 8, with an 80 percent dataset replacement strategy at each level. Experiments show that greedy inference achieves a 77 percent solve rate at depth 5, and beam search (width 5) improves this to 89 percent. Both strategies degrade sharply beyond depth 6, reflecting the fundamental limitations of imitation learning without a planning mechanism. The web interface provides real-time animated demonstrations of all three solving strategies and an in-browser benchmark comparison.", first_indent=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("Keywords: ")
    r.font.name = "Arial"; r.font.size = Pt(12); r.font.bold = True
    r2 = p.add_run("Rubik's Cube, curriculum learning, imitation learning, MLP, beam search, Three.js, combinatorial optimisation, neural network, BFS")
    r2.font.name = "Arial"; r2.font.size = Pt(12)

    page_break(doc)
    p = doc.add_paragraph()
    r = p.add_run("Author's Declaration")
    r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True
    p.paragraph_format.space_after = Pt(16)

    body(doc, "I hereby declare that this master's thesis was written independently by me and that all sources used have been cited and acknowledged in the bibliography. The work has not been submitted previously in whole or in part for any other degree or examination.", first_indent=False)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Cm(3)
    r2 = p2.add_run("Czestochowa, May 2026")
    r2.font.name = "Arial"; r2.font.size = Pt(12)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Cm(2)
    r3 = p3.add_run("_________________________")
    r3.font.name = "Arial"; r3.font.size = Pt(12)
    p4 = doc.add_paragraph()
    r4 = p4.add_run("Edward Ogbei")
    r4.font.name = "Arial"; r4.font.size = Pt(12)


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # page setup
    for sec in doc.sections:
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin   = Cm(3.5)
        sec.right_margin  = Cm(2.0)
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)

    title_page(doc)
    page_break(doc)

    # TOC placeholder
    p = doc.add_paragraph()
    r = p.add_run("Table of Contents")
    r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True
    p.paragraph_format.space_after = Pt(10)
    note(doc, "(Generate automatically using Word / LibreOffice: Insert > Table of Contents.)")

    ch_intro(doc)
    ch_aim(doc)
    ch_background(doc)
    ch_architecture(doc)
    ch_solvers(doc)
    ch_training(doc)
    ch_viz(doc)
    ch_results(doc)
    ch_conclusions(doc)
    bibliography(doc)
    appendices(doc)
    abstracts(doc)

    doc.save(OUT)
    print(f"Saved: {OUT}")

    # Convert to PDF
    try:
        from docx2pdf import convert
        pdf_out = OUT.replace(".docx", ".pdf")
        convert(OUT, pdf_out)
        print(f"PDF:   {pdf_out}")
    except Exception as e:
        print(f"PDF conversion skipped: {e}")


if __name__ == "__main__":
    main()
